# -*- coding: utf-8 -*-
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor

INSTR = Path(
    r"C:\Users\estua\Downloads\Capitulo_IV_Ingenieria_de_Requerimientos\instrumentos"
)
SIG = INSTR / "firma_sello_solo.png"


def set_run(run, size=11, bold=False):
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(size)
    run.bold = bold
    run.font.color.rgb = RGBColor(0, 0, 0)


def clear(p):
    for c in list(p._p):
        if c.tag.endswith("}r") or c.tag.endswith("}hyperlink"):
            p._p.remove(c)


def main():
    doc_path = INSTR / "Guia_Entrevista_Requerimientos_SIGEC_IGSS_FIRMADA.docx"
    if not doc_path.exists():
        doc_path = INSTR / "Guia_Entrevista_Requerimientos_SIGEC_IGSS_LLENADA.docx"
    doc = Document(str(doc_path))

    for p in doc.paragraphs:
        if "Firma del participante" not in p.text:
            continue
        clear(p)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.first_line_indent = Cm(0)
        r = p.add_run("Firma del participante (firma digital):\n")
        set_run(r, bold=True)
        img = p.add_run()
        img.add_picture(str(SIG), width=Inches(2.8))
        p.add_run("\n")
        for line in [
            "Lic. Luis Gustavo Sierra González",
            'Administrador "C"',
            "Consultorio de Palín, Escuintla",
        ]:
            rr = p.add_run(line + "\n")
            set_run(rr, size=10)
        rr = p.add_run("\nFirma del entrevistador(a): _______________________________")
        set_run(rr)
        break

    out = INSTR / "Guia_Entrevista_Requerimientos_SIGEC_IGSS_FIRMADA.docx"
    try:
        doc.save(str(out))
    except PermissionError:
        out = INSTR / "Guia_Entrevista_Requerimientos_SIGEC_IGSS_FIRMADA_v2.docx"
        doc.save(str(out))
    print(out)


if __name__ == "__main__":
    main()
