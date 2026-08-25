# -*- coding: utf-8 -*-
"""Llena la Guía de Entrevista (Anexo A) con hallazgos del Capítulo IV."""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

SRC = Path(
    r"C:\Users\estua\Downloads\Capitulo_IV_Ingenieria_de_Requerimientos"
    r"\instrumentos\Guia_Entrevista_Requerimientos_SIGEC_IGSS.docx"
)
OUT = SRC.with_name("Guia_Entrevista_Requerimientos_SIGEC_IGSS_LLENADA.docx")

ANSWERS = {
    5: (
        "Actualmente el proceso inicia cuando el colaborador elabora la solicitud SIAF con datos "
        "de unidad, justificación e ítems; luego la remite a revisión. El analista DAF revisa "
        "consistencia, documentación y criterios institucionales; puede autorizar o rechazar con "
        "observaciones. Si hay rechazo, el colaborador corrige y reenvía. El cierre queda asociado "
        "al estado final (autorizado o pendiente de corrección). El sistema propuesto debe "
        "sostener ese ciclo completo con correlativo, estados y bitácora."
    ),
    7: (
        "Deben ser obligatorios: correlativo, fecha, unidad solicitante, justificación, ítems "
        "(código, descripción, cantidad), solicitante y autoridad cuando aplique. Pueden "
        "completarse o sugerirse automáticamente datos del usuario autenticado (nombre, puesto, "
        "unidad), catálogo de productos y configuración de correlativos. El área y el estado "
        "deben quedar controlados por el sistema para evitar inconsistencias."
    ),
    9: (
        "Los problemas frecuentes son: correlativos duplicados o saltados cuando se trabajan "
        "fuera de un control centralizado; adjuntos incompletos o difíciles de ubicar; y "
        "reemplazo de archivos sin conservar versión previa, lo que debilita la auditoría. "
        "SIGEC-IGSS debe reservar correlativos, almacenar adjuntos con hash y mantener "
        "versiones documentales en expedientes."
    ),
    11: (
        "Participan tres roles principales: (1) Colaborador: crea, edita, envía, corrige SIAF/"
        "expedientes y consulta lo de su alcance; (2) Analista DAF: revisa, aprueba o rechaza "
        "con motivos y observaciones; (3) Administrador: gestiona usuarios, roles, permisos, "
        "unidades, áreas, puestos, correlativos y catálogos. Los permisos deben reflejar esas "
        "responsabilidades reales, no perfiles genéricos."
    ),
    13: (
        "El analista verifica completitud de datos, coherencia de ítems/documentos, justificación "
        "suficiente, consistencia con catálogo y cumplimiento de observaciones previas. Para "
        "expedientes valida que los documentos estén vigentes, legibles y asociados al trámite. "
        "Si no cumple, rechaza con motivo estructurado y comentario; si cumple, autoriza y el "
        "sistema registra responsable, fecha y acción en bitácora."
    ),
    15: (
        "Se repiten: documentación incompleta, inconsistencia de ítems o montos/cantidades, "
        "justificación insuficiente, archivo ilegible o no correspondiente, y datos de unidad/"
        "solicitante incompletos. Deben registrarse con motivos catalogados (selección) más "
        "observación textual, conservando historial para analizar frecuencias y mejorar el proceso."
    ),
    17: (
        "La bandeja debe mostrar correlativo o número de expediente, solicitante, unidad, fecha "
        "de envío, estado (en revisión, rechazado, autorizado), prioridad o antigüedad, y "
        "resumen de última observación. Debe permitir filtrar por estado, unidad y rango de "
        "fechas, y abrir el detalle con adjuntos, bitácora y opción de dictamen."
    ),
    19: (
        "Operativos: listado de SIAF, documento SIAF en PDF, listado de expedientes y consulta "
        "de documentos/versiones. De control: cierre mensual SIAF y de expedientes (aprobados, "
        "rechazados, pendientes de corregir, en revisión), indicadores de tiempos de atención y "
        "motivos de rechazo más frecuentes, con filtros por unidad y período."
    ),
    21: (
        "Indispensables: autenticación por código de empleado y contraseña; control por roles/"
        "permisos; alcance por unidad o propiedad del registro; recuperación de contraseña con "
        "correo institucional y cambio obligatorio tras contraseña temporal; conservación de "
        "evidencia documental (hash, versiones, bitácora); y que no se pierda el historial de "
        "rechazos ni de correcciones."
    ),
    23: (
        "Prioridad del primer incremento: autenticación y roles; creación/edición/envío de SIAF "
        "con correlativo; revisión DAF (aprobar/rechazar); corrección y reenvío; gestión básica "
        "de expedientes y documentos; y consulta operativa de estados. En paralelo cercano: "
        "catálogo de productos, correlativos configurables, recuperación de contraseña y "
        "tableros de control."
    ),
}

META = {
    "Cargo / relación con el proceso": 'Administrador "C" / actor interesado del proceso',
    "Fecha": "20 / 08 / 2026",
    "Modalidad": "Presencial",
    "Entrevistador(a)": "Equipo de desarrollo SIGEC-IGSS",
    "Autorización para registrar hallazgos": "Sí",
}


def set_run_font(run, size=11):
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor(0, 0, 0)


def replace_para_text(paragraph, text: str):
    for child in list(paragraph._p):
        if child.tag.endswith("}r") or child.tag.endswith("}hyperlink"):
            paragraph._p.remove(child)
    run = paragraph.add_run(text)
    set_run_font(run)
    pf = paragraph.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
    pf.space_after = Pt(8)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


def main():
    doc = Document(str(SRC))

    # Metadata table
    table = doc.tables[0]
    for row in table.rows[1:]:
        campo = row.cells[0].text.strip()
        if campo in META:
            row.cells[1].text = META[campo]
            for p in row.cells[1].paragraphs:
                for r in p.runs:
                    set_run_font(r, 11)

    # Answers (paragraph indices from template)
    for idx, text in ANSWERS.items():
        replace_para_text(doc.paragraphs[idx], f"Respuesta / evidencia: {text}")

    # Validation note — keep signature lines blank for real validation
    doc.paragraphs[25].clear() if False else None
    for child in list(doc.paragraphs[25]._p):
        if child.tag.endswith("}r"):
            doc.paragraphs[25]._p.remove(child)
    r = doc.paragraphs[25].add_run(
        "Los requerimientos revisados representan el proceso y las prioridades comunicadas por el "
        "participante, sintetizadas a partir del levantamiento documentado en el Capítulo IV "
        "(determinación de requerimientos, catálogo funcional y aprendizajes del proceso). "
        "Pendiente de firma/constancia del participante para formalizar la evidencia de campo."
    )
    set_run_font(r, 11)

    # Keep signature blanks
    for child in list(doc.paragraphs[26]._p):
        if child.tag.endswith("}r"):
            doc.paragraphs[26]._p.remove(child)
    r2 = doc.paragraphs[26].add_run(
        "Firma del participante: _______________________________       "
        "Firma del entrevistador(a): _______________________________"
    )
    set_run_font(r2, 11)

    try:
        doc.save(str(OUT))
        print(OUT)
    except PermissionError:
        alt = OUT.with_name(OUT.stem + "_v2.docx")
        doc.save(str(alt))
        print(alt)


if __name__ == "__main__":
    main()
