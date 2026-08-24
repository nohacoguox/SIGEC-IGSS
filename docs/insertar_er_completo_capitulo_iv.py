# -*- coding: utf-8 -*-
"""Inserta el modelo ER completo (5 partes) en el Capítulo IV con formato APA 7."""

from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor
from docx.text.paragraph import Paragraph

DOC = Path(
    r"C:\Users\estua\Downloads\Capitulo_IV_Ingenieria_de_Requerimientos"
    r"\Capitulo_IV_SIGEC_IGSS_FINAL_CON_RECUPERACION.docx"
)
OUT_ALT = DOC.with_name("Capitulo_IV_SIGEC_IGSS_FINAL_CON_ER_COMPLETO.docx")
IMG_DIR = Path(r"c:\PROYECTOS-PERSONALES\SIGEC-IGSS\docs\diagramas_er")

PARTS = [
    {
        "file": "er_parte_a_mapa_modulos.png",
        "title": "Modelo entidad–relación completo de SIGEC-IGSS — Parte A: mapa de módulos",
        "note": (
            "Se presenta la organización del modelo en cinco módulos lógicos. Las partes B–E "
            "detallan atributos, claves y cardinalidades. Elaboración propia a partir de las "
            "entidades TypeORM del prototipo (25 entidades y 2 tablas de unión)."
        ),
    },
    {
        "file": "er_parte_b_seguridad_org.png",
        "title": "Modelo entidad–relación — Parte B: seguridad, acceso y estructura organizacional",
        "note": (
            "Incluye Usuario, Credencial, Rol, Permiso, tablas de unión user_roles y "
            "role_permissions, Puesto, Departamento, Municipio, Unidad médica y Área. "
            "La relación Usuario–Unidad médica es lógica (texto), no FK. Elaboración propia."
        ),
    },
    {
        "file": "er_parte_c_siaf.png",
        "title": "Modelo entidad–relación — Parte C: módulo SIAF",
        "note": (
            "Muestra SIAF_SOLICITUD y sus dependencias 1:N (ítems, subproductos, autorizaciones, "
            "bitácora y documentos adjuntos), además de vínculos con Usuario y Área. "
            "El vínculo con ProductoCatalogo es lógico por código y origen. Elaboración propia."
        ),
    },
    {
        "file": "er_parte_d_expedientes.png",
        "title": "Modelo entidad–relación — Parte D: módulo de expedientes",
        "note": (
            "Incluye Expediente, documentos, versiones, bitácora y detalle de observaciones "
            "documentales. El campo numero_siaf vincula lógicamente el expediente con una "
            "solicitud SIAF sin FK ORM. Elaboración propia."
        ),
    },
    {
        "file": "er_parte_e_catalogos_correlativos.png",
        "title": "Modelo entidad–relación — Parte E: catálogo de productos y correlativos",
        "note": (
            "Cubre ProductoCatalogo, ProductoCatalogoConfig, configuración y reserva de "
            "correlativos SIAF, y configuración de correlativos de expedientes. Elaboración propia."
        ),
    },
]

INTRO = (
    "El modelo entidad–relación resume las entidades persistentes que soportan las reglas de "
    "negocio de SIGEC-IGSS. Dado el volumen del esquema (seguridad, estructura organizacional, "
    "SIAF, expedientes, catálogos y correlativos), el modelo completo se presenta en cinco partes "
    "secuenciales (Figuras 16 a 20), a fin de conservar legibilidad sin omitir entidades. "
    "Las líneas discontinuas indican asociaciones lógicas sin clave foránea en el ORM; las "
    "continuas representan relaciones implementadas con integridad referencial."
)


def set_run(run, *, size=12, bold=False, italic=False):
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = RGBColor(0, 0, 0)


def clear_para(p):
    for child in list(p._p):
        if child.tag.endswith("}r") or child.tag.endswith("}hyperlink"):
            p._p.remove(child)


def format_body(p):
    pf = p.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.first_line_indent = Cm(1.27)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    for r in p.runs:
        set_run(r)


def add_para_after(anchor: Paragraph, text: str = "") -> Paragraph:
    el = OxmlElement("w:p")
    anchor._p.addnext(el)
    return Paragraph(el, anchor._parent)


def write_fig_caption(p: Paragraph, num: int):
    clear_para(p)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf = p.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    pf.space_before = Pt(12)
    pf.space_after = Pt(0)
    pf.first_line_indent = Cm(0)
    r = p.add_run(f"Figura {num}")
    set_run(r, bold=True)


def write_fig_title(p: Paragraph, title: str):
    clear_para(p)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf = p.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    pf.space_before = Pt(0)
    pf.space_after = Pt(6)
    pf.first_line_indent = Cm(0)
    r = p.add_run(title)
    set_run(r, italic=True)


def write_note(p: Paragraph, note: str):
    clear_para(p)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf = p.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    pf.space_before = Pt(0)
    pf.space_after = Pt(12)
    pf.first_line_indent = Cm(0)
    r1 = p.add_run("Nota. ")
    set_run(r1, italic=True)
    r2 = p.add_run(note)
    set_run(r2)


def para_has_drawing(p: Paragraph) -> bool:
    return bool(p._p.xpath(".//*[local-name()='drawing' or local-name()='pict']"))


def remove_old_er_block(doc: Document):
    """Elimina desde la intro antigua / Figura 16 hasta antes de 4.4 o Conclusión."""
    paras = doc.paragraphs
    start = None
    end = None
    for i, p in enumerate(paras):
        t = p.text.strip()
        if t.startswith("4.3.3"):
            # keep heading; start removing after it
            start = i + 1
        if start is not None and i > start and (
            t.startswith("4.4 ") or t == "Conclusión del capítulo" or t.startswith("4.4\t")
        ):
            end = i
            break
    if start is None:
        raise SystemExit("No se encontró 4.3.3")
    if end is None:
        end = len(paras)

    # Collect elements to remove between start and end
    to_remove = []
    for i in range(start, end):
        to_remove.append(paras[i]._p)
    for el in to_remove:
        parent = el.getparent()
        if parent is not None:
            parent.remove(el)
    return start


def insert_er_section(doc: Document):
    # Find heading 4.3.3
    heading = None
    for p in doc.paragraphs:
        if p.text.strip().startswith("4.3.3"):
            heading = p
            break
    if heading is None:
        raise SystemExit("Heading 4.3.3 no encontrado tras limpieza")

    # Determine next figure number: max existing Figura N before 4.3.3, then continue
    max_fig = 15
    for p in doc.paragraphs:
        t = p.text.strip()
        m = re.match(r"^Figura\s+(\d+)$", t)
        if m:
            max_fig = max(max_fig, int(m.group(1)))
        if p._p is heading._p:
            break
    # After removing old fig 16, max should be 15; start at 16
    first_num = max_fig + 1 if max_fig >= 15 else 16

    cursor = heading

    # Intro
    intro_p = add_para_after(cursor, "")
    clear_para(intro_p)
    r = intro_p.add_run(INTRO)
    set_run(r)
    format_body(intro_p)
    cursor = intro_p

    for idx, part in enumerate(PARTS):
        num = first_num + idx
        img_path = IMG_DIR / part["file"]
        if not img_path.exists():
            raise SystemExit(f"Falta imagen: {img_path}")

        # Caption number
        cap = add_para_after(cursor)
        write_fig_caption(cap, num)
        cursor = cap

        # Title
        title_p = add_para_after(cursor)
        write_fig_title(title_p, part["title"])
        cursor = title_p

        # Image
        img_p = add_para_after(cursor)
        clear_para(img_p)
        img_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        img_p.paragraph_format.first_line_indent = Cm(0)
        img_p.paragraph_format.space_before = Pt(0)
        img_p.paragraph_format.space_after = Pt(6)
        run = img_p.add_run()
        run.add_picture(str(img_path), width=Inches(6.5))
        cursor = img_p

        # Note
        note_p = add_para_after(cursor)
        write_note(note_p, part["note"])
        cursor = note_p

    return first_num, first_num + len(PARTS) - 1


def main():
    doc = Document(str(DOC))
    remove_old_er_block(doc)
    first, last = insert_er_section(doc)
    try:
        doc.save(str(DOC))
        saved = DOC
    except PermissionError:
        doc.save(str(OUT_ALT))
        saved = OUT_ALT
        print("AVISO: documento original abierto; se guardó copia.")
    print(f"Figuras {first}–{last} insertadas.")
    print(saved)


if __name__ == "__main__":
    main()
