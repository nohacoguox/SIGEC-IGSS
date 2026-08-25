# -*- coding: utf-8 -*-
from pathlib import Path
from docx import Document

path = Path(
    r"c:\Users\estua\OneDrive\Documents\UMG\10MO\seminario\INFORMES\INFORME II\PROPUESTA NOE Y GEOVANY.docx"
)

MEMBERS = [
    ("Walter Francisco Sontay Vicente", "5390-14-11522"),
    ("Noe Estuardo Coguox Mach", "5390-17-3396"),
    ("Jarod Fernando Fernandez Morales", "5390-19-15689"),
    ("Héctor Hugo Barrera Contreras", "5390-20-12591"),
    ("Juan Carlos Bautista Catota", "5390-20-1601"),
    ("Geovany Emmanuel González Díaz", "5390-22-848"),
]

lines = ["Integrantes del grupo (Sección A):"]
for n, (name, carnet) in enumerate(MEMBERS, 1):
    lines.append(f"{n}. {name}    Carné: {carnet}")


def set_para_text(p, text):
    for r in p.runs:
        r.text = ""
    if p.runs:
        p.runs[0].text = text
    else:
        p.add_run(text)


doc = Document(str(path))
paras = doc.paragraphs

proyecto_idx = next(i for i, p in enumerate(paras) if "Proyecto de seminario" in p.text)
origen_idx = next(i for i, p in enumerate(paras) if p.text.strip().startswith("1. Origen y fundación"))

# Remove paragraphs between proyecto and origen (exclusive)
for p in paras[proyecto_idx + 1 : origen_idx]:
    p._element.getparent().remove(p._element)

# Reload paragraphs after removal
doc.save(str(path))
doc = Document(str(path))
paras = doc.paragraphs
proyecto_idx = next(i for i, p in enumerate(paras) if "Proyecto de seminario" in p.text)
origen_p = next(p for p in paras if p.text.strip().startswith("1. Origen y fundación"))

for line in reversed(lines):
    origen_p.insert_paragraph_before(line)

doc.save(str(path))

# verify
doc2 = Document(str(path))
for i, p in enumerate(doc2.paragraphs[:14]):
    print(i, p.text[:85] if p.text else "(empty)")
