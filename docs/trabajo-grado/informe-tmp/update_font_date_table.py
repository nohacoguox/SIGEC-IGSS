# -*- coding: utf-8 -*-
"""Unify fonts, add delivery date, put members in table on proposal."""
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_ALIGN_PARAGRAPH

BASE = Path(r"c:\Users\estua\OneDrive\Documents\UMG\10MO\seminario\INFORMES\INFORME II")

MEMBERS = [
    ("Walter Francisco Sontay Vicente", "5390-14-11522"),
    ("Noe Estuardo Coguox Mach", "5390-17-3396"),
    ("Jarod Fernando Fernandez Morales", "5390-19-15689"),
    ("Héctor Hugo Barrera Contreras", "5390-20-12591"),
    ("Juan Carlos Bautista Catota", "5390-20-1601"),
    ("Geovany Emmanuel González Díaz", "5390-22-848"),
]

FECHA_ENTREGA = "6 de agosto de 2026"
FECHA_LINE = f"Guatemala, {FECHA_ENTREGA}"
FONT_NAME = "Times New Roman"
FONT_SIZE = Pt(11)

FILES_ALL = [
    "INFORME_II_Institucion_Elegida_EORM_Agua_de_la_Mina_ACTUALIZADO.docx",
    "borrador_INFORME_II_Institucion_Elegida_EORM_Agua_de_la_Mina.docx",
    "PROPUESTA_FORMAL_EORM_Agua_de_la_Mina.docx",
    "SOLICITUD_INFORMACION_EORM_Agua_de_la_Mina.docx",
    "PROPUESTA NOE Y GEOVANY.docx",
]


def set_run_font(run, size=None, bold=None):
    run.font.name = FONT_NAME
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:ascii"), FONT_NAME)
    rFonts.set(qn("w:hAnsi"), FONT_NAME)
    rFonts.set(qn("w:eastAsia"), FONT_NAME)
    rFonts.set(qn("w:cs"), FONT_NAME)
    if size is not None:
        run.font.size = size
    if bold is not None:
        run.bold = bold


def unify_fonts(doc):
    # Default style
    style = doc.styles["Normal"]
    style.font.name = FONT_NAME
    style.font.size = FONT_SIZE
    rPr = style.element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:ascii"), FONT_NAME)
    rFonts.set(qn("w:hAnsi"), FONT_NAME)
    rFonts.set(qn("w:eastAsia"), FONT_NAME)
    rFonts.set(qn("w:cs"), FONT_NAME)

    for p in doc.paragraphs:
        for run in p.runs:
            # keep existing size if set, otherwise 11
            size = run.font.size if run.font.size else FONT_SIZE
            set_run_font(run, size=size)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for run in p.runs:
                        size = run.font.size if run.font.size else FONT_SIZE
                        set_run_font(run, size=size)

    # Also headings if present
    for sname in doc.styles:
        try:
            st = doc.styles[sname]
            if hasattr(st, "font"):
                st.font.name = FONT_NAME
                rPr = st.element.get_or_add_rPr()
                rFonts = rPr.find(qn("w:rFonts"))
                if rFonts is None:
                    rFonts = OxmlElement("w:rFonts")
                    rPr.insert(0, rFonts)
                rFonts.set(qn("w:ascii"), FONT_NAME)
                rFonts.set(qn("w:hAnsi"), FONT_NAME)
                rFonts.set(qn("w:eastAsia"), FONT_NAME)
                rFonts.set(qn("w:cs"), FONT_NAME)
        except Exception:
            pass


def set_para_text(p, text):
    for r in p.runs:
        r.text = ""
    if p.runs:
        p.runs[0].text = text
        set_run_font(p.runs[0], size=FONT_SIZE)
    else:
        run = p.add_run(text)
        set_run_font(run, size=FONT_SIZE)


def shade_header_cell(cell, hex_color="1F4E79"):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_color)
    shd.set(qn("w:val"), "clear")
    tcPr.append(shd)
    for p in cell.paragraphs:
        for run in p.runs:
            run.font.color.rgb = RGBColor(255, 255, 255)
            set_run_font(run, size=Pt(11), bold=True)


def set_cell(cell, text, bold=False, size=11, white=False):
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(text)
    set_run_font(run, size=Pt(size), bold=bold)
    if white:
        run.font.color.rgb = RGBColor(255, 255, 255)


def add_date_near_top(doc, after_markers):
    """Insert fecha de entrega if not present."""
    full = "\n".join(p.text for p in doc.paragraphs)
    if FECHA_ENTREGA in full or "Fecha de entrega" in full:
        # update existing if placeholder
        for p in doc.paragraphs:
            if "Fecha de entrega" in p.text and FECHA_ENTREGA not in p.text:
                set_para_text(p, f"Fecha de entrega: {FECHA_ENTREGA}")
                return "updated"
        return "exists"

    # Find insertion point after institution header / subtitle
    insert_after = None
    for i, p in enumerate(doc.paragraphs):
        for m in after_markers:
            if m in p.text:
                insert_after = i
    if insert_after is None:
        insert_after = min(8, len(doc.paragraphs) - 1)

    # Insert before the paragraph after insert_after
    if insert_after + 1 < len(doc.paragraphs):
        anchor = doc.paragraphs[insert_after + 1]
    else:
        anchor = doc.add_paragraph()

    # Insert blank + date lines before anchor (reversed)
    lines = [
        f"Fecha de entrega: {FECHA_ENTREGA}",
        FECHA_LINE,
    ]
    for line in reversed(lines):
        np = anchor.insert_paragraph_before(line)
        for r in np.runs:
            set_run_font(r, size=FONT_SIZE)
        if not np.runs:
            r = np.add_run(line)
            set_run_font(r, size=FONT_SIZE)
    return "inserted"


def replace_integrantes_with_table(doc):
    """Replace numbered integrante paragraphs with a Word table in the proposal."""
    # Find header and member lines
    start_idx = None
    member_idxs = []
    for i, p in enumerate(doc.paragraphs):
        t = p.text.strip()
        if "Integrantes del grupo" in t or t.startswith("Integrantes"):
            start_idx = i
        if start_idx is not None and i > start_idx:
            if t[:2].rstrip(".").isdigit() and "Carné" in t:
                member_idxs.append(i)
            elif t.startswith("ACEPTACIÓN") or t.startswith("Yo,") or t.startswith("Recibido"):
                break
            elif member_idxs and t and not t[0].isdigit():
                break

    if start_idx is None or not member_idxs:
        return False

    # Anchor = paragraph after last member (for table insertion before it)
    last = member_idxs[-1]
    if last + 1 < len(doc.paragraphs):
        anchor = doc.paragraphs[last + 1]
    else:
        anchor = doc.add_paragraph()

    # Create table by inserting before anchor using document.add_table then moving XML
    # Simpler approach: clear member paragraphs, put table after header using insert
    header_p = doc.paragraphs[start_idx]
    set_para_text(header_p, "Integrantes del grupo — Sección A")

    # Remove member paragraphs
    for i in reversed(member_idxs):
        el = doc.paragraphs[i]._element
        el.getparent().remove(el)

    # Reload references: find header again
    header_p = None
    for p in doc.paragraphs:
        if "Integrantes del grupo" in p.text:
            header_p = p
            break
    if header_p is None:
        return False

    # Find next paragraph after header to insert table before it
    found = False
    next_p = None
    for p in doc.paragraphs:
        if found:
            next_p = p
            break
        if p._element is header_p._element:
            found = True
    if next_p is None:
        next_p = doc.add_paragraph()

    # Add table at end then move before next_p
    table = doc.add_table(rows=1 + len(MEMBERS), cols=3)
    table.style = "Table Grid"
    set_cell(table.rows[0].cells[0], "No.", bold=True, white=True)
    set_cell(table.rows[0].cells[1], "Nombre completo", bold=True, white=True)
    set_cell(table.rows[0].cells[2], "Carné", bold=True, white=True)
    for c in table.rows[0].cells:
        shade_header_cell(c)

    for i, (name, carnet) in enumerate(MEMBERS, start=1):
        set_cell(table.rows[i].cells[0], str(i))
        set_cell(table.rows[i].cells[1], name)
        set_cell(table.rows[i].cells[2], carnet)

    # Move table XML before next_p
    tbl = table._tbl
    next_p._element.addprevious(tbl)

    # Remove empty leftover table placeholder if add_table left it at end - already moved
    return True


def process_all():
    for fname in FILES_ALL:
        path = BASE / fname
        if not path.exists():
            print("SKIP", fname)
            continue
        doc = Document(str(path))

        # 1) Unify fonts
        unify_fonts(doc)

        # 2) Date for propuesta and solicitud
        if fname in (
            "PROPUESTA_FORMAL_EORM_Agua_de_la_Mina.docx",
            "SOLICITUD_INFORMACION_EORM_Agua_de_la_Mina.docx",
        ):
            markers = [
                "Directora:",
                "EORM Agua de la Mina — Jornada Vespertina",
                "Aldea Agua de la Mina",
                "control de asistencia estudiantil",
            ]
            result = add_date_near_top(doc, markers)
            print(f"  date: {result}")

            # Also ensure Atentamente block mentions fecha if useful
            # Add under Atentamente? Better near top is enough.
            # Fill blank date lines near signature if they are delivery-related
            for p in doc.paragraphs:
                if p.text.strip().startswith("Lugar y fecha:") and "agosto" not in p.text:
                    # Keep institutional acceptance blank; delivery date is separate
                    pass

        # 3) Table only for propuesta formal
        if fname == "PROPUESTA_FORMAL_EORM_Agua_de_la_Mina.docx":
            ok = replace_integrantes_with_table(doc)
            print(f"  table: {ok}")
            unify_fonts(doc)  # re-apply after table creation

        doc.save(str(path))
        print(f"UPDATED {fname}")


if __name__ == "__main__":
    process_all()
