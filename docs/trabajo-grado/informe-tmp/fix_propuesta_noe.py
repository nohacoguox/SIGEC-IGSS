# -*- coding: utf-8 -*-
"""Fix integrantes order in PROPUESTA NOE Y GEOVANY.docx"""
from pathlib import Path
from docx import Document

path = Path(
    r"c:\Users\estua\OneDrive\Documents\UMG\10MO\seminario\INFORMES\INFORME II\PROPUESTA NOE Y GEOVANY.docx"
)

MEMBERS = [
    ("Walter Francisco Sontay Vicente", "5390-14-11522"),
    ("Noe Estuardo Coguox Mach", "5390-17-3396"),
    ("Jarod Fernando Fernandez Morales", "5390-19-15689"),
    ("Héctor Hugo Barrera Contreras", "5390-20-12591"),
    ("Juan Carlos Bautista Catota", "5390-20-1601"),
    ("Geovany Emmanuel González Díaz", "5390-22-848"),
]

doc = Document(str(path))

# Remove misplaced integrante lines and header
to_remove = []
for i, p in enumerate(doc.paragraphs):
    t = p.text.strip()
    if t == "Integrantes del grupo (Sección A):":
        to_remove.append(p)
    elif t.startswith("1. Walter") or t.startswith("2. Noe") or t.startswith("3. Jarod"):
        to_remove.append(p)
    elif t.startswith("4. Héctor") or t.startswith("5. Juan") or t.startswith("6. Geovany"):
        to_remove.append(p)
    elif t == "" and i < 15:
        # empty para near top - might be from insert; check neighbors
        pass

for p in to_remove:
    el = p._element
    el.getparent().remove(el)

# Find proyecto paragraph and insert after it
ref = None
for p in doc.paragraphs:
    if "Proyecto de seminario" in p.text:
        ref = p
        break

if ref is None:
    raise SystemExit("No proyecto paragraph found")

# Insert after ref: add paragraphs in order using insert_paragraph_after pattern
# python-docx only has insert_paragraph_before on paragraph; insert after by using next para
lines = ["Integrantes del grupo (Sección A):"]
for n, (name, carnet) in enumerate(MEMBERS, 1):
    lines.append(f"{n}. {name}    Carné: {carnet}")

# insert before the paragraph after ref
next_p = None
found = False
for p in doc.paragraphs:
    if found:
        next_p = p
        break
    if p is ref:
        found = True

if next_p is None:
    next_p = doc.add_paragraph()

for line in reversed(lines):
    next_p.insert_paragraph_before(line)

doc.save(str(path))
print("Fixed PROPUESTA NOE Y GEOVANY.docx")

# verify
doc2 = Document(str(path))
for p in doc2.paragraphs[:12]:
    if p.text.strip():
        print(p.text)
