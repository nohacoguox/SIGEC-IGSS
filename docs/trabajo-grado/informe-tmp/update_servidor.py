# -*- coding: utf-8 -*-
"""Update docs to include server + development + implementation + biometrics with installation."""
from pathlib import Path
from docx import Document

BASE = Path(r"c:\Users\estua\OneDrive\Documents\UMG\10MO\seminario\INFORMES\INFORME II")

FILES = [
    "INFORME_II_Institucion_Elegida_EORM_Agua_de_la_Mina_ACTUALIZADO.docx",
    "borrador_INFORME_II_Institucion_Elegida_EORM_Agua_de_la_Mina.docx",
    "PROPUESTA_FORMAL_EORM_Agua_de_la_Mina.docx",
    "SOLICITUD_INFORMACION_EORM_Agua_de_la_Mina.docx",
    "PROPUESTA NOE Y GEOVANY.docx",
]


def set_para_text(p, text):
    for r in p.runs:
        r.text = ""
    if p.runs:
        p.runs[0].text = text
    else:
        p.add_run(text)


def replace_in_doc(doc, old, new):
    count = 0
    for p in doc.paragraphs:
        if old in p.text:
            set_para_text(p, p.text.replace(old, new))
            count += 1
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    if old in p.text:
                        set_para_text(p, p.text.replace(old, new))
                        count += 1
    return count


# Exact / near-exact replacements applied across docs
REPLACEMENTS = [
    # Subtítulo propuesta formal
    (
        "Donación e instalación de dos (2) relojes biométricos e implementación de software\n"
        "de control de asistencia estudiantil",
        "Donación de servidor, desarrollo e implementación de software, e instalación de dos (2)\n"
        "relojes biométricos para el control de asistencia estudiantil",
    ),
    (
        "Donación e instalación de dos (2) relojes biométricos e implementación de software de control de asistencia estudiantil",
        "Donación de servidor, desarrollo e implementación de software, e instalación de dos (2) "
        "relojes biométricos para el control de asistencia estudiantil",
    ),
    # Proyecto NOE / GEO
    (
        "Proyecto de seminario:\xa0Donar dispositivos biométricos e implementar software de control de asistencia "
        "(reconocimiento facial, huella dactilar y/o código del alumno), porque actualmente la asistencia se lleva\xa0"
        "manual en cuadernillos.\xa0No cuentan con biométricos en operación hoy.",
        "Proyecto de seminario:\xa0Brindar un servidor; desarrollar e implementar el software de control de asistencia; "
        "y donar e instalar dos (2) relojes biométricos (reconocimiento facial, huella dactilar y/o código del alumno), "
        "porque actualmente la asistencia se lleva\xa0manual en cuadernillos.\xa0No cuentan con biométricos en operación hoy.",
    ),
    (
        "Proyecto seminario en evaluación: donación de biométricos + software de asistencia.",
        "Proyecto seminario en evaluación: servidor + desarrollo e implementación de software + donación e instalación de biométricos.",
    ),
    (
        "Proyecto de seminario en evaluación: donación de biométricos + software de asistencia.",
        "Proyecto de seminario en evaluación: servidor + desarrollo e implementación de software + donación e instalación de biométricos.",
    ),
    # Informe actualizado - sección VI intro
    (
        "El grupo de seminario propone donar e instalar dos (2) relojes biométricos e implementar un software "
        "de control de asistencia estudiantil, como aporte tecnológico orientado a resolver el cuello de botella "
        "del registro manual. Se establecen dos dispositivos por la cantidad aproximada de 320 alumnos, a fin de "
        "evitar aglomeraciones en el ingreso. El alcance del aporte se limita a lo propuesto: puesta en funcionamiento "
        "de los biométricos y del software asociado; no incluye mejoras adicionales fuera de dicho alcance.",
        "El grupo de seminario propone brindar un servidor, desarrollar e implementar un software de control de "
        "asistencia estudiantil, y donar e instalar dos (2) relojes biométricos, como aporte tecnológico orientado "
        "a resolver el cuello de botella del registro manual. Se establecen dos dispositivos por la cantidad "
        "aproximada de 320 alumnos, a fin de evitar aglomeraciones en el ingreso. El alcance del aporte se limita "
        "a: servidor, desarrollo, implementación del software e instalación/puesta en marcha de los biométricos; "
        "no incluye mejoras adicionales fuera de dicho alcance.",
    ),
    # Borrador original (texto más corto)
    (
        "El grupo de seminario propone donar dispositivos biométricos e implementar un software de control de "
        "asistencia estudiantil, como aporte tecnológico de bajo costo orientado a resolver el cuello de botella "
        "del registro manual.",
        "El grupo de seminario propone brindar un servidor, desarrollar e implementar un software de control de "
        "asistencia estudiantil, y donar e instalar dos (2) relojes biométricos, como aporte tecnológico orientado "
        "a resolver el cuello de botella del registro manual. El alcance se limita a: servidor, desarrollo, "
        "implementación e instalación de biométricos.",
    ),
    # Objetivo propuesta formal
    (
        "Digitalizar y automatizar el control de asistencia estudiantil de la EORM Agua de la Mina — JV, "
        "mediante la donación e instalación de dos relojes biométricos y un software de gestión, "
        "mejorando la trazabilidad, oportunidad y confiabilidad de la información para dirección, "
        "docentes y familias.",
        "Digitalizar y automatizar el control de asistencia estudiantil de la EORM Agua de la Mina — JV, "
        "mediante la donación de un servidor, el desarrollo e implementación de un software de gestión, "
        "y la donación e instalación de dos (2) relojes biométricos, mejorando la trazabilidad, oportunidad "
        "y confiabilidad de la información para dirección, docentes y familias.",
    ),
    # Alcance bullets / items
    (
        "Dos (2) relojes biométricos (reconocimiento facial y/o lector de huella), con opción "
        "de código del alumno como respaldo, según factibilidad técnica, ética y de consentimiento.",
        "Servidor para alojar y operar el sistema de asistencia.",
    ),
    (
        "Dispositivos biométricos (reconocimiento facial y/o lector de huella), según "
        "factibilidad técnica, ética y de consentimiento.",
        "Servidor para alojar y operar el sistema de asistencia.",
    ),
    (
        "Software de control de asistencia: registro diario, consultas por alumno/grado/fecha, consolidación y reportes.",
        "Desarrollo e implementación de software de control de asistencia: registro diario, consultas por alumno/grado/fecha, consolidación y reportes.",
    ),
    (
        "Donar e instalar dos (2) relojes biométricos. Se establecen dos equipos por el volumen "
        "aproximado de 320 alumnos, para reducir aglomeraciones en el ingreso y facilitar el flujo.",
        "Brindar un servidor para el sistema de control de asistencia.",
    ),
    (
        "Dejar los biométricos funcionando, incluyendo la gestión de las conexiones eléctricas y "
        "de red necesarias para su operación (aprovechando que el instituto ya cuenta con energía, "
        "internet y data center).",
        "Desarrollar e implementar el software de control de asistencia (registro, consulta y reportes).",
    ),
    (
        "Desarrollar e implementar un software de control de asistencia (registro, consulta y reportes).",
        "Donar e instalar dos (2) relojes biométricos, dejándolos funcionando, incluyendo las conexiones "
        "eléctricas y de red necesarias (aprovechando que el instituto ya cuenta con energía, internet "
        "y data center). Se establecen dos equipos por el volumen aproximado de 320 alumnos, para reducir "
        "aglomeraciones en el ingreso.",
    ),
    # Fuera de alcance: quitar "servidores" del fuera de alcance porque SÍ lo brindan
    (
        "Compra de computadoras adicionales, servidores o ampliación de infraestructura no relacionada.",
        "Compra de computadoras adicionales o ampliación de infraestructura no relacionada "
        "(distinta al servidor y biométricos del aporte).",
    ),
    # Conclusiones
    (
        "La solución propuesta —donación e instalación de dos (2) relojes biométricos e implementación "
        "de software de control de asistencia— responde de forma directa a ese requerimiento, con "
        "alcance limitado a lo ofrecido (sin mejoras adicionales fuera de la propuesta).",
        "La solución propuesta —servidor, desarrollo e implementación de software, y donación e instalación "
        "de dos (2) relojes biométricos— responde de forma directa a ese requerimiento, con alcance limitado "
        "a lo ofrecido (sin mejoras adicionales fuera de la propuesta).",
    ),
    (
        "La solución propuesta —donación de dispositivos biométricos e implementación de software "
        "de control de asistencia— responde de forma directa a ese requerimiento, siempre bajo "
        "consideraciones éticas de protección de datos de menores y validación técnica in situ.",
        "La solución propuesta —servidor, desarrollo e implementación de software, y donación e instalación "
        "de biométricos— responde de forma directa a ese requerimiento, siempre bajo consideraciones éticas "
        "de protección de datos de menores y validación técnica in situ.",
    ),
    # Solicitud / referencias cortas
    (
        "aporte tecnológico (dos relojes biométricos y software de control de asistencia estudiantil)",
        "aporte tecnológico (servidor, desarrollo e implementación de software, e instalación de dos relojes biométricos)",
    ),
    (
        "Entregar propuesta formal al instituto (2 relojes biométricos + software) y checklist "
        "de ubicación física / responsabilidades compartidas.",
        "Entregar propuesta formal al instituto (servidor + desarrollo/implementación + 2 relojes biométricos "
        "con instalación) y checklist de ubicación física / responsabilidades compartidas.",
    ),
    (
        "se explicó el aporte tecnológico (relojes biométricos + software de asistencia)",
        "se explicó el aporte tecnológico (servidor + desarrollo/implementación de software + relojes biométricos con instalación)",
    ),
    (
        "Explicación del aporte: biométricos + software de asistencia.",
        "Explicación del aporte: servidor + desarrollo e implementación de software + biométricos con instalación.",
    ),
    (
        "biométricos + software de asistencia",
        "servidor + desarrollo/implementación de software + biométricos con instalación",
    ),
    (
        "propuesta formal (2 relojes biométricos + software)",
        "propuesta formal (servidor + desarrollo/implementación + 2 biométricos con instalación)",
    ),
    # Tabla comparación anexo si existe
    (
        "Biométricos + software de asistencia",
        "Servidor + desarrollo/implementación + biométricos con instalación",
    ),
]


def insert_bullets_after_match(doc, match_substr, new_bullets):
    """If match becomes a single-item replacement that removed content, add missing bullets after it."""
    for i, p in enumerate(doc.paragraphs):
        if match_substr in p.text:
            # insert after this paragraph using next paragraph as anchor
            if i + 1 < len(doc.paragraphs):
                anchor = doc.paragraphs[i + 1]
            else:
                anchor = doc.add_paragraph()
            for bullet in reversed(new_bullets):
                np = anchor.insert_paragraph_before(bullet)
                try:
                    np.style = "List Bullet"
                except Exception:
                    pass
            return True
    return False


def ensure_informe_alcance(doc):
    """Ensure section 6.2 has full scope bullets for informe docs."""
    # After 'Servidor para alojar...' add the missing items if not already present
    full = "\n".join(p.text for p in doc.paragraphs)
    if "Dos (2) relojes biométricos (reconocimiento facial" in full and "Servidor para alojar" in full:
        return False
    if "Servidor para alojar y operar el sistema de asistencia." not in full:
        return False

    extra = []
    if "Dos (2) relojes biométricos (reconocimiento facial" not in full:
        extra.append(
            "Dos (2) relojes biométricos con instalación (reconocimiento facial y/o lector de huella), "
            "con opción de código del alumno como respaldo, según factibilidad técnica, ética y de consentimiento."
        )
    if "Modalidad alternativa o complementaria: código del alumno" not in full:
        # may already exist; skip if present
        pass
    if extra:
        return insert_bullets_after_match(
            doc,
            "Servidor para alojar y operar el sistema de asistencia.",
            extra,
        )
    return False


def ensure_propuesta_alcance(doc):
    """Ensure propuesta formal has server + development + biometrics bullets."""
    full = "\n".join(p.text for p in doc.paragraphs)
    # After the first commitment bullets sequence, verify all three pillars exist
    needed = []
    if "Brindar un servidor" not in full and "Servidor para alojar" not in full:
        needed.append("Brindar un servidor para el sistema de control de asistencia.")
    # development may already be there via replacements
    if "Desarrollar e implementar el software" not in full and "Desarrollar e implementar un software" not in full:
        needed.append("Desarrollar e implementar el software de control de asistencia (registro, consulta y reportes).")
    if "Donar e instalar dos (2) relojes biométricos, dejándolos funcionando" not in full:
        if "Donar e instalar dos (2) relojes biométricos." not in full:
            needed.append(
                "Donar e instalar dos (2) relojes biométricos, dejándolos funcionando, incluyendo las conexiones "
                "eléctricas y de red necesarias."
            )
    return needed


def process(path: Path):
    doc = Document(str(path))
    total = 0
    for old, new in REPLACEMENTS:
        n = replace_in_doc(doc, old, new)
        total += n

    # Informe: ensure biometric bullet remains after server bullet
    ensure_informe_alcance(doc)

    # Propuesta formal: if software bullet was overwritten awkwardly, fix sequence
    if path.name.startswith("PROPUESTA_FORMAL"):
        # Re-check and patch "El grupo se compromete a" section if incomplete
        texts = [p.text for p in doc.paragraphs]
        joined = "\n".join(texts)
        # Fix subtitle if still old
        for p in doc.paragraphs:
            if p.text.strip().startswith("Donación e instalación de dos"):
                set_para_text(
                    p,
                    "Donación de servidor, desarrollo e implementación de software, e instalación de dos (2) "
                    "relojes biométricos para el control de asistencia estudiantil",
                )
                total += 1

        # Ensure modalidad / capacitación bullets still make sense after reshuffle
        # Find "El grupo se compromete a:" and rewrite next commitment bullets cleanly
        for i, p in enumerate(doc.paragraphs):
            if p.text.strip() == "El grupo se compromete a:":
                # Collect following bullets until "Queda fuera del alcance"
                j = i + 1
                commit_paras = []
                while j < len(doc.paragraphs):
                    t = doc.paragraphs[j].text.strip()
                    if t.startswith("Queda fuera del alcance"):
                        break
                    if t:
                        commit_paras.append(doc.paragraphs[j])
                    j += 1

                desired = [
                    "Brindar un servidor para alojar y operar el sistema de control de asistencia.",
                    "Desarrollar e implementar el software de control de asistencia (registro, consulta y reportes).",
                    "Donar e instalar dos (2) relojes biométricos, dejándolos funcionando, incluyendo las conexiones "
                    "eléctricas y de red necesarias (aprovechando que el instituto ya cuenta con energía, internet "
                    "y data center). Se establecen dos equipos por el volumen aproximado de 320 alumnos, para reducir "
                    "aglomeraciones en el ingreso.",
                    "Habilitar modalidades de marcaje según factibilidad y autorización institucional: "
                    "reconocimiento facial, huella dactilar y/o código del alumno (como respaldo o alternativa).",
                    "Capacitar al personal enlace designado por la dirección.",
                    "Entregar documentación básica de uso y recomendaciones de resguardo de la información.",
                ]
                # Overwrite existing non-empty commit paras; if fewer, insert; if more keep extras only if useful
                for idx, text in enumerate(desired):
                    if idx < len(commit_paras):
                        set_para_text(commit_paras[idx], text)
                    else:
                        # insert before "Queda fuera..."
                        fuera = None
                        for q in doc.paragraphs:
                            if q.text.strip().startswith("Queda fuera del alcance"):
                                fuera = q
                                break
                        if fuera:
                            fuera.insert_paragraph_before(text)
                total += 1
                break

        # Update plan table phase 3 if present
        for table in doc.tables:
            for row in table.rows:
                cells = [c.text for c in row.cells]
                if cells and "3. Implementación" in cells[0]:
                    if len(row.cells) > 1:
                        set_para_text(
                            row.cells[1].paragraphs[0],
                            "Instalación del servidor; despliegue del software; instalación eléctrica/red de los "
                            "biométricos; pruebas de operación.",
                        )
                        total += 1
                if cells and "2. Preparación" in cells[0]:
                    if len(row.cells) > 1 and "servidor" not in row.cells[1].text.lower():
                        set_para_text(
                            row.cells[1].paragraphs[0],
                            "Reunión con padres; consentimientos; relevamiento técnico de puntos de instalación "
                            "de biométricos y ubicación del servidor.",
                        )
                        total += 1

    # Informe 6.2: if first bullet became only server, restore biometric + software wording
    if "INFORME" in path.name.upper() or path.name.startswith("borrador_"):
        for i, p in enumerate(doc.paragraphs):
            if p.text.strip() == "Servidor para alojar y operar el sistema de asistencia.":
                # Check next paragraphs
                nxt = doc.paragraphs[i + 1].text if i + 1 < len(doc.paragraphs) else ""
                if "Dos (2) relojes" not in nxt and "relojes biométricos" not in nxt:
                    # Insert biometric and ensure software line mentions desarrollo/implementación
                    anchor = doc.paragraphs[i + 1] if i + 1 < len(doc.paragraphs) else doc.add_paragraph()
                    np = anchor.insert_paragraph_before(
                        "Dos (2) relojes biométricos con instalación (reconocimiento facial y/o lector de huella), "
                        "con opción de código del alumno como respaldo, según factibilidad técnica, ética y de consentimiento."
                    )
                    try:
                        np.style = p.style
                    except Exception:
                        pass
                    total += 1
                break

        # Beneficios / table cells mentioning Ausencia de biométricos solution
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if "Donación e instalación de dispositivos, más software de gestión asociado." in cell.text:
                        replace_in_doc(
                            doc,
                            "Donación e instalación de dispositivos, más software de gestión asociado.",
                            "Donación de servidor, desarrollo/implementación de software e instalación de biométricos.",
                        )
                        total += 1

    doc.save(str(path))
    return total


def main():
    for fname in FILES:
        path = BASE / fname
        if not path.exists():
            print("SKIP", fname)
            continue
        n = process(path)
        print(f"UPDATED {fname} ({n} replacements/patches)")


if __name__ == "__main__":
    main()
