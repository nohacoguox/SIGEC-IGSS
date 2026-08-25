# -*- coding: utf-8 -*-
from pathlib import Path
from docx import Document

path = Path(
    r"c:\Users\estua\OneDrive\Documents\UMG\10MO\seminario\INFORMES\INFORME II\PROPUESTA NOE Y GEOVANY.docx"
)

MEMBERS = [
    ("Walter Francisco Sontay Vicente", "5390-14-11522"),
    ("Noe Estuardo Coguox Mach", "5390-17-3396"),
    ("Jarod Fernando Fernandez Morales", "5390-19-15689"),
    ("Héctor Hugo Barrera Contreras", "5390-20-12591"),
    ("Juan Carlos Bautista Catota", "5390-20-1601"),
    ("Geovany Emmanuel González Díaz", "5390-22-848"),
]

doc = Document(str(path))

# Remove empty paragraphs at indices 4-5 if still empty
paras = doc.paragraphs
proyecto_idx = None
for i, p in enumerate(paras):
    if "Proyecto de seminario" in p.text:
        proyecto_idx = i
        break

if proyecto_idx is None:
    raise SystemExit("no proyecto")

# Check if integrantes already correct after proyecto
if proyecto_idx + 1 < len(paras) and "Integrantes del grupo" in paras[proyecto_idx + 1].text:
    print("Already OK")
else:
    # Insert after proyecto using document body manipulation
    ref = paras[proyecto_idx]
    lines = ["Integrantes del grupo (Sección A):"]
    for n, (name, carnet) in enumerate(MEMBERS, 1):
        lines.append(f"{n}. {name}    Carné: {carnet}")

    # insert_paragraph_before on paragraph after proyecto (or before section 1)
    anchor = paras[proyecto_idx + 1] if proyecto_idx + 1 < len(paras) else None
    if anchor is None:
        anchor = doc.add_paragraph()

    for line in lines:
        new_p = anchor.insert_paragraph_before(line)
        anchor = new_p  # chain inserts in order

    doc.save(str(path))
    print("Inserted integrantes")

# verify
doc2 = Document(str(path))
for i, p in enumerate(doc2.paragraphs[:15]):
    print(i, repr(p.text[:85]) if p.text else "(empty)")
