# -*- coding: utf-8 -*-
from pathlib import Path
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

path = Path(
    r"c:\Users\estua\OneDrive\Documents\UMG\10MO\seminario\INFORMES\INFORME II\PROPUESTA_FORMAL_EORM_Agua_de_la_Mina.docx"
)
FONT = "Times New Roman"


def set_run_font(run, size=Pt(11)):
    run.font.name = FONT
    run.font.size = size
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:ascii"), FONT)
    rFonts.set(qn("w:hAnsi"), FONT)
    rFonts.set(qn("w:eastAsia"), FONT)
    rFonts.set(qn("w:cs"), FONT)


doc = Document(str(path))

# Remove misplaced date lines near the end
for p in list(doc.paragraphs):
    t = p.text.strip()
    if t in ("Guatemala, 6 de agosto de 2026", "Fecha de entrega: 6 de agosto de 2026"):
        p._element.getparent().remove(p._element)

# Insert after Directora line
anchor = None
for p in doc.paragraphs:
    if p.text.strip().startswith("Directora:"):
        # next paragraph after this one
        found = False
        for q in doc.paragraphs:
            if found:
                anchor = q
                break
            if q._element is p._element:
                found = True
        break

if anchor is None:
    raise SystemExit("Directora not found")

for line in reversed(["Guatemala, 6 de agosto de 2026", "Fecha de entrega: 6 de agosto de 2026"]):
    np = anchor.insert_paragraph_before(line)
    if np.runs:
        set_run_font(np.runs[0])
    else:
        r = np.add_run(line)
        set_run_font(r)

doc.save(str(path))

# verify
d = Document(str(path))
print("TOP:")
for i, p in enumerate(d.paragraphs[:14]):
    print(i, p.text[:90])
print("Any leftover dates at end?")
for i, p in enumerate(d.paragraphs):
    if "Fecha de entrega" in p.text or p.text.strip().startswith("Guatemala, 6"):
        print(i, p.text)
