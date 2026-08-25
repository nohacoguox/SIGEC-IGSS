# -*- coding: utf-8 -*-
"""Update existing INFORME II documents with team member info."""
from pathlib import Path
from docx import Document

BASE = Path(r"c:\Users\estua\OneDrive\Documents\UMG\10MO\seminario\INFORMES\INFORME II")

MEMBERS = [
    ("Walter Francisco Sontay Vicente", "5390-14-11522"),
    ("Noe Estuardo Coguox Mach", "5390-17-3396"),
    ("Jarod Fernando Fernandez Morales", "5390-19-15689"),
    ("Héctor Hugo Barrera Contreras", "5390-20-12591"),
    ("Juan Carlos Bautista Catota", "5390-20-1601"),
    ("Geovany Emmanuel González Díaz", "5390-22-848"),
]

SECCION = "A"
NAMES_LIST = ", ".join(m[0] for m in MEMBERS)
NAMES_LIST_Y = (
    "Walter Francisco Sontay Vicente, Noe Estuardo Coguox Mach, "
    "Jarod Fernando Fernandez Morales, Héctor Hugo Barrera Contreras, "
    "Juan Carlos Bautista Catota y Geovany Emmanuel González Díaz"
)

FILES = [
    "INFORME_II_Institucion_Elegida_EORM_Agua_de_la_Mina_ACTUALIZADO.docx",
    "borrador_INFORME_II_Institucion_Elegida_EORM_Agua_de_la_Mina.docx",
    "SOLICITUD_INFORMACION_EORM_Agua_de_la_Mina.docx",
    "PROPUESTA_FORMAL_EORM_Agua_de_la_Mina.docx",
]


def set_para_text(p, text):
    for r in p.runs:
        r.text = ""
    if p.runs:
        p.runs[0].text = text
    else:
        p.add_run(text)


def update_integrantes_paragraphs(doc):
    """Replace numbered integrante placeholder lines."""
    updated = 0
    for i, p in enumerate(doc.paragraphs):
        t = p.text.strip()
        for n, (name, carnet) in enumerate(MEMBERS, start=1):
            # Long format (informe)
            if t.startswith(f"{n}.") and "Carné" in t and ("___" in t or not any(m[0] in t for m in MEMBERS)):
                set_para_text(p, f"{n}. {name}    Carné: {carnet}")
                updated += 1
                break
    return updated


def update_completar_note(doc):
    for p in doc.paragraphs:
        if "Completar nombre completo" in p.text:
            set_para_text(p, f"Sección: {SECCION}")
            return True
    return False


def update_table_participantes(doc):
    updated = 0
    for table in doc.tables:
        for row in table.rows:
            if not row.cells:
                continue
            label = row.cells[0].text.strip()
            if label == "Participantes del grupo" and len(row.cells) > 1:
                val = row.cells[1].text.strip()
                if "___" in val or "Completar" in val or val.startswith("[") or not val:
                    row.cells[1].text = ""
                    row.cells[1].paragraphs[0].add_run(NAMES_LIST_Y)
                    updated += 1
    return updated


def replace_in_all(doc, old, new):
    count = 0
    for p in doc.paragraphs:
        if old in p.text:
            set_para_text(p, p.text.replace(old, new))
            count += 1
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    if old in p.text:
                        set_para_text(p, p.text.replace(old, new))
                        count += 1
    return count


def process_file(path: Path):
    doc = Document(str(path))
    stats = {
        "integrantes": update_integrantes_paragraphs(doc),
        "seccion": update_completar_note(doc),
        "participantes_table": update_table_participantes(doc),
    }

  # Text replacements for examples / placeholders
    replacements = [
        ("[Completar nombres de los integrantes presentes el 31/07/2026]", NAMES_LIST_Y),
        ("[Completar nombres exactos y pegar evidencia fotográfica.]", NAMES_LIST_Y + ". [Pegar evidencia fotográfica.]"),
        ("Participantes: [nombres del grupo]", f"Participantes: {NAMES_LIST_Y}"),
        ("Participantes del grupo: [nombres]", f"Participantes del grupo: {NAMES_LIST_Y}"),
    ]
    for old, new in replacements:
        stats[f"repl_{old[:20]}"] = replace_in_all(doc, old, new)

    # Fig 3 example line update
    for p in doc.paragraphs:
        if "viernes 31/07/2026" in p.text and "Completar nombres" in p.text:
            set_para_text(
                p,
                p.text.replace(
                    "[Completar nombres exactos y pegar evidencia fotográfica.]",
                    NAMES_LIST_Y + ". [Pegar evidencia fotográfica.]",
                ),
            )

    doc.save(str(path))
    return stats


def main():
    for fname in FILES:
        path = BASE / fname
        if not path.exists():
            print(f"SKIP (not found): {fname}")
            continue
        stats = process_file(path)
        print(f"UPDATED: {fname}")
        print(f"  stats: {stats}")


if __name__ == "__main__":
    main()
