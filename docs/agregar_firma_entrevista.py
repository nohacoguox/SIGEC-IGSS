# -*- coding: utf-8 -*-
"""Extrae la firma y la inserta en la guía de entrevista como firma digital."""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor
from PIL import Image

ASSET = Path(
    r"C:\Users\estua\.cursor\projects\c-PROYECTOS-PERSONALES-SIGEC-IGSS\assets"
    r"\c__Users_estua_AppData_Roaming_Cursor_User_workspaceStorage_"
    r"067ede5504a840b3d43caca8d74fd55c_images_image-70484933-538f-4449-aa7b-59428e151cf6.png"
)
INSTR = Path(
    r"C:\Users\estua\Downloads\Capitulo_IV_Ingenieria_de_Requerimientos\instrumentos"
)
SIG_PNG = INSTR / "firma_digital_luis_gustavo_sierra.png"
DOC = INSTR / "Guia_Entrevista_Requerimientos_SIGEC_IGSS_LLENADA.docx"
OUT = INSTR / "Guia_Entrevista_Requerimientos_SIGEC_IGSS_LLENADA.docx"


def prepare_signature(src: Path, dst: Path) -> Path:
    im = Image.open(src).convert("RGBA")
    w, h = im.size
    # Recortar "Atentamente," (aprox. 18% superior)
    top = int(h * 0.16)
    im = im.crop((0, top, w, h))

    pixels = im.load()
    for y in range(im.height):
        for x in range(im.width):
            r, g, b, a = pixels[x, y]
            # Fondo blanco/casi blanco → transparente
            if r > 235 and g > 235 and b > 235:
                pixels[x, y] = (255, 255, 255, 0)
            else:
                pixels[x, y] = (r, g, b, 255)

    # Recorte al contenido no transparente
    bbox = im.getbbox()
    if bbox:
        im = im.crop(bbox)

    # Lienzo limpio con margen
    pad = 12
    canvas = Image.new("RGBA", (im.width + pad * 2, im.height + pad * 2), (255, 255, 255, 0))
    canvas.paste(im, (pad, pad), im)
    canvas.save(dst, "PNG")
    return dst


def set_run(run, size=11, bold=False, italic=False):
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = RGBColor(0, 0, 0)


def clear_para(p):
    for child in list(p._p):
        if child.tag.endswith("}r") or child.tag.endswith("}hyperlink"):
            p._p.remove(child)


def main():
    prepare_signature(ASSET, SIG_PNG)

    doc = Document(str(DOC))

    # Actualizar validación: ya no "pendiente de firma"
    for p in doc.paragraphs:
        t = p.text.strip()
        if t.startswith("Los requerimientos revisados"):
            clear_para(p)
            r = p.add_run(
                "Los requerimientos revisados representan el proceso y las prioridades "
                "comunicadas por el participante. Validado con firma digital del actor "
                "interesado."
            )
            set_run(r)
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

        if t.startswith("Firma del participante"):
            clear_para(p)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.first_line_indent = Cm(0)
            p.paragraph_format.space_before = Pt(6)

            # Bloque participante (izquierda)
            label = p.add_run("Firma del participante (firma digital):\n")
            set_run(label, bold=True, size=11)
            run_img = p.add_run()
            run_img.add_picture(str(SIG_PNG), width=Inches(2.6))
            p.add_run("\n")
            id_line = p.add_run(
                'Lic. Luis Gustavo Sierra González\nAdministrador "C"\n'
                "Consultorio de Palín, Escuintla\n"
            )
            set_run(id_line, size=10)
            sep = p.add_run("\nFirma del entrevistador(a): _______________________________")
            set_run(sep, size=11)

    # Asegurar cargo en tabla
    table = doc.tables[0]
    for row in table.rows:
        if row.cells[0].text.strip().startswith("Cargo"):
            row.cells[1].text = 'Administrador "C" — Consultorio de Palín, Escuintla'

    try:
        doc.save(str(OUT))
        saved = OUT
    except PermissionError:
        saved = OUT.with_name("Guia_Entrevista_Requerimientos_SIGEC_IGSS_FIRMADA.docx")
        doc.save(str(saved))

    print(SIG_PNG)
    print(saved)


if __name__ == "__main__":
    main()
