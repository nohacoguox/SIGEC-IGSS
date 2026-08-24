#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Genera el documento Word APA del recorrido de prototipo SIGEC-IGSS."""

from __future__ import annotations

import json
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


OUT_DIR = Path.home() / "Downloads" / "Recorrido_SIGEC_IGSS"
CAPTURAS = OUT_DIR / "capturas"
MANIFEST = CAPTURAS / "manifest.json"
DOC_PATH = OUT_DIR / "Recorrido_Prototipo_SIGEC_IGSS_APA.docx"

MESES = {
    1: "enero", 2: "febrero", 3: "marzo", 4: "abril", 5: "mayo", 6: "junio",
    7: "julio", 8: "agosto", 9: "septiembre", 10: "octubre", 11: "noviembre", 12: "diciembre",
}

EXTRA_TEXT = {
    "Inicio de sesión": (
        "Esta interfaz constituye el punto de ingreso controlado al sistema. El usuario autentica su identidad "
        "mediante código de empleado y contraseña, lo que habilita la asignación dinámica de paneles y pantallas "
        "según el perfil institucional autorizado."
    ),
    "Panel de administración": (
        "Luego del acceso privilegiado, el administrador visualiza indicadores agregados y dispone de un menú "
        "lateral permanente para desplazarse entre dashboard, reportes, configuración y gestiones maestras."
    ),
}


def set_run_font(run, name="Times New Roman", size=12, bold=False, italic=False):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = RGBColor(0, 0, 0)


def add_para(doc, text, *, bold=False, italic=False, size=12, align="justify", space_after=8, first_line=True):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    pf.space_after = Pt(space_after)
    pf.space_before = Pt(0)
    if first_line and align == "justify":
        pf.first_line_indent = Cm(1.27)
    if align == "center":
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pf.first_line_indent = Cm(0)
    elif align == "right":
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        pf.first_line_indent = Cm(0)
    elif align == "left":
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        pf.first_line_indent = Cm(0)
    else:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold, italic=italic)
    return p


def add_heading_apa(doc, text, level=1):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    pf.space_before = Pt(12)
    pf.space_after = Pt(12)
    pf.first_line_indent = Cm(0)
    if level == 1:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        set_run_font(run, size=12, bold=True)
    elif level == 2:
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(text)
        set_run_font(run, size=12, bold=True)
    else:
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(text)
        set_run_font(run, size=12, bold=True, italic=True)
    return p


def add_figure(doc, image_path: Path, figura_n: int, titulo: str, nota: str):
    p_img = doc.add_paragraph()
    p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_img.paragraph_format.first_line_indent = Cm(0)
    p_img.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p_img.paragraph_format.space_before = Pt(6)
    p_img.paragraph_format.space_after = Pt(6)
    run = p_img.add_run()
    run.add_picture(str(image_path), width=Cm(15.2))

    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.LEFT
    cap.paragraph_format.first_line_indent = Cm(0)
    cap.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    cap.paragraph_format.space_before = Pt(6)
    r1 = cap.add_run(f"Figura {figura_n}")
    set_run_font(r1, bold=True)
    cap.add_run("\n")
    r2 = cap.add_run(titulo)
    set_run_font(r2, italic=True)

    nota_p = doc.add_paragraph()
    nota_p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    nota_p.paragraph_format.first_line_indent = Cm(0)
    nota_p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    rn = nota_p.add_run(f"Nota. {nota}")
    set_run_font(rn, size=10)


def fecha_es():
    now = datetime.now()
    return f"{now.day} de {MESES[now.month]} de {now.year}"


def build():
    if not MANIFEST.exists():
        raise SystemExit(f"No se encontró {MANIFEST}")

    items = json.loads(MANIFEST.read_text(encoding="utf-8"))
    OUT_DIR.mkdir(parents=True, exist_ok=True)

    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)

    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")

    # Portada APA
    for _ in range(3):
        add_para(doc, "", align="center", first_line=False, space_after=0)
    add_para(
        doc,
        "Recorrido funcional del prototipo del Sistema Integral de Gestión de Expedientes de Compras (SIGEC-IGSS): "
        "pantallas, navegación e interacción de usuario",
        align="center",
        bold=True,
        first_line=False,
        space_after=24,
    )
    add_para(doc, "Estuardo [Apellido del autor]", align="center", first_line=False, space_after=0)
    add_para(doc, "Instituto Guatemalteco de Seguridad Social (IGSS)", align="center", first_line=False, space_after=0)
    add_para(doc, "[Nombre del curso / programa académico]", align="center", first_line=False, space_after=0)
    add_para(doc, "[Nombre del docente]", align="center", first_line=False, space_after=0)
    add_para(doc, fecha_es(), align="center", first_line=False, space_after=0)

    doc.add_page_break()

    add_heading_apa(doc, "Resumen", 1)
    add_para(
        doc,
        "El presente documento describe el recorrido funcional del prototipo operativo SIGEC-IGSS, desarrollado como "
        "avance real del sistema institucional de gestión de solicitudes SIAF y expedientes de compras. Se documentan "
        f"las pantallas principales ({len(items)} capturas), la navegación entre módulos y la interacción del usuario "
        "con los flujos de autenticación, administración, operación, revisión DAF y analítica. Las capturas "
        "corresponden al estado actual del software y se presentan conforme a las recomendaciones de la séptima "
        "edición del Manual de publicaciones de la American Psychological Association (APA, 2020) para figuras e "
        "informes técnicos.",
    )
    add_para(
        doc,
        "Palabras clave: SIGEC-IGSS, prototipo funcional, SIAF, expedientes de compras, navegación de usuario, APA.",
        italic=True,
        first_line=False,
    )

    add_heading_apa(doc, "Introducción", 1)
    add_para(
        doc,
        "SIGEC-IGSS es una plataforma web orientada a digitalizar y controlar el ciclo de vida de las solicitudes SIAF "
        "y de los expedientes de compras del Instituto Guatemalteco de Seguridad Social. A diferencia de un mockup "
        "estático producido únicamente en herramientas de prototipado (por ejemplo, Moqups), el presente recorrido se "
        "construye sobre avances reales del sistema: autenticación por roles, paneles administrativos y operativos, "
        "libro SIAF, expedientes, bandeja de revisiones DAF y tableros analíticos. El objetivo es evidenciar, con "
        "capturas secuenciales, la navegación e interacción disponibles para la evaluación académica del prototipo "
        "funcional y, cuando se requiera, complementar con un video demostrativo.",
    )

    add_heading_apa(doc, "Método", 1)
    add_heading_apa(doc, "Materiales y entorno", 2)
    add_para(
        doc,
        "Las capturas se obtuvieron del frontend React del sistema SIGEC-IGSS en ejecución local, con resolución de "
        "referencia de 1440 × 900 píxeles. Cada figura incluye título, descripción de la interacción y nota "
        "aclaratoria sobre el propósito funcional de la pantalla dentro del flujo institucional.",
    )
    add_heading_apa(doc, "Procedimiento de recorrido", 2)
    add_para(
        doc,
        "El recorrido sigue el orden natural de uso: (1) acceso y autenticación; (2) panel administrativo y gestiones "
        "maestras; (3) panel colaborador y operación SIAF/expedientes; (4) revisión DAF; (5) analítica SIAF y de "
        "expedientes; y (6) seguridad de credenciales, incluyendo recuperación y cambio de contraseña. En cada paso "
        "se documenta qué observa el usuario, qué acciones puede ejecutar y cómo avanza hacia la siguiente pantalla.",
    )
    add_heading_apa(doc, "Mapa de navegación", 2)
    add_para(
        doc,
        "Login (¿Olvidaste tu contraseña? → solicitud de recuperación → correo institucional) → "
        "(elección de panel, si aplica) → Panel administración (Gestiones: Usuarios, Roles, Áreas, Puestos, "
        "Unidades Médicas, Correlativos) ↔ Panel colaborador → Libro SIAF (Listado / Crear) → Expedientes de compras → "
        "Bandeja de revisiones DAF (SIAF / Expedientes) → Estadísticas (Análisis SIAF / Análisis de expedientes) → "
        "Cambio de contraseña / Cierre de sesión.",
        first_line=True,
    )

    add_heading_apa(doc, "Resultados: recorrido de pantallas", 1)
    add_para(
        doc,
        f"Se documentaron {len(items)} capturas del prototipo. A continuación se presenta la secuencia completa con su "
        "explicación profesional orientada a la evaluación de navegación e interacción.",
    )

    current_section = None
    figura_n = 1
    for item in items:
        seccion = item.get("seccion") or "General"
        if seccion != current_section:
            current_section = seccion
            add_heading_apa(doc, seccion, 2)

        add_heading_apa(doc, item["titulo"], 3)
        texto = item.get("explicacion") or ""
        extra = EXTRA_TEXT.get(item["titulo"])
        add_para(doc, texto if not extra else f"{texto} {extra}")

        img = Path(item["file"])
        if not img.is_file():
            img = CAPTURAS / Path(item["file"]).name
        if img.is_file():
            add_figure(
                doc,
                img,
                figura_n,
                item["titulo"],
                "Captura de pantalla del prototipo funcional SIGEC-IGSS. La figura ilustra la interfaz disponible "
                "para la navegación e interacción del usuario en este punto del flujo operativo.",
            )
            figura_n += 1
        else:
            add_para(doc, f"[No se encontró la captura para {item['titulo']}]", italic=True, first_line=False)

    add_heading_apa(doc, "Discusión", 1)
    add_para(
        doc,
        "El recorrido confirma que el sistema dispone de una cobertura funcional cercana a la operación completa: "
        "autenticación, configuración institucional, captura y seguimiento SIAF, gestión documental de expedientes, "
        "revisión DAF y análisis de resultados. Esta evidencia permite presentar avances reales como prototipo "
        "navegable, reduciendo la dependencia de herramientas de mockup para las pantallas ya implementadas y "
        "reservando prototipado complementario solo para eventuales pendientes menores. En consecuencia, el "
        "entregable satisface el requerimiento de visualizar pantallas, navegación e interacción de usuario con base "
        "en software desarrollado, no únicamente en wireframes.",
    )

    add_heading_apa(doc, "Conclusión", 1)
    add_para(
        doc,
        "SIGEC-IGSS, en su estado actual, constituye un prototipo funcional demostrable mediante recorrido de "
        "pantallas reales. Las figuras presentadas documentan la navegación e interacción del usuario de extremo a "
        "extremo y pueden utilizarse como evidencia académica y técnica del avance del proyecto, así como apoyo "
        "visual para un video demostrativo del sistema.",
    )

    add_heading_apa(doc, "Referencias", 1)
    refs = [
        "American Psychological Association. (2020). Publication manual of the American Psychological Association (7.ª ed.). https://doi.org/10.1037/0000165-000",
        "Instituto Guatemalteco de Seguridad Social. (s. f.). Marco institucional de gestión administrativa y financiera.",
        "SIGEC-IGSS. (2026). Sistema Integral de Gestión de Expedientes de Compras [Software en desarrollo].",
    ]
    for ref in refs:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
        p.paragraph_format.first_line_indent = Cm(-1.27)
        p.paragraph_format.left_indent = Cm(1.27)
        run = p.add_run(ref)
        set_run_font(run)

    doc.save(DOC_PATH)
    print(f"Documento generado: {DOC_PATH}")
    print(f"Figuras: {figura_n - 1}")
    print(f"Carpeta: {OUT_DIR}")


if __name__ == "__main__":
    build()
