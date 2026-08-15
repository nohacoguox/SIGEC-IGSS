# -*- coding: utf-8 -*-
"""Add team info to PROPUESTA NOE Y GEOVANY.docx"""
from pathlib import Path
from docx import Document

path = Path(
    r"c:\Users\estua\OneDrive\Documents\UMG\10MO\seminario\INFORMES\INFORME II\PROPUESTA NOE Y GEOVANY.docx"
)

MEMBERS = [
    ("Walter Francisco Sontay Vicente", "5390-14-11522"),
    ("Noe Estuardo Coguox Mach", "5390-17-3396"),
    ("Jarod Fernando Fernandez Morales", "5390-19-15689"),
    ("Héctor Hugo Barrera Contreras", "5390-20-12591"),
    ("Juan Carlos Bautista Catota", "5390-20-1601"),
    ("Geovany Emmanuel González Díaz", "5390-22-848"),
]

doc = Document(str(path))

# Skip if already has team
full = "\n".join(p.text for p in doc.paragraphs)
if "Walter Francisco Sontay Vicente" in full and "5390-14-11522" in full:
    print("Already has team info")
else:
    # Find insertion point after proyecto line (paragraph 3)
    insert_idx = None
    for i, p in enumerate(doc.paragraphs):
        if "Proyecto de seminario" in p.text:
            insert_idx = i + 1
            break

    if insert_idx is None:
        insert_idx = 4

    # Build new paragraphs to insert - use element after reference para
    ref = doc.paragraphs[min(insert_idx, len(doc.paragraphs)-1)]
    lines = [
        "",
        "Integrantes del grupo (Sección A):",
    ]
    for n, (name, carnet) in enumerate(MEMBERS, 1):
        lines.append(f"{n}. {name}    Carné: {carnet}")

    for line in reversed(lines):
        new_p = ref.insert_paragraph_before(line)

    doc.save(str(path))
    print("Updated PROPUESTA NOE Y GEOVANY.docx")
