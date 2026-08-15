# -*- coding: utf-8 -*-
from pathlib import Path
from docx import Document

path = Path(
    r"c:\Users\estua\OneDrive\Documents\UMG\10MO\seminario\INFORMES\INFORME II\PROPUESTA NOE Y GEOVANY.docx"
)

MEMBERS = [
    ("Walter Francisco Sontay Vicente", "5390-14-11522"),
    ("Noe Estuardo Coguox Mach", "5390-17-3396"),
    ("Jarod Fernando Fernandez Morales", "5390-19-15689"),
    ("Héctor Hugo Barrera Contreras", "5390-20-12591"),
    ("Juan Carlos Bautista Catota", "5390-20-1601"),
    ("Geovany Emmanuel González Díaz", "5390-22-848"),
]

doc = Document(str(path))

# Remove bad integrante block
remove_texts = {
    "Integrantes del grupo (Sección A):",
}
for p in list(doc.paragraphs):
    t = p.text.strip()
    if t in remove_texts or (t.startswith(tuple(f"{i}." for i in range(1, 7))) and "Carné:" in t):
        p._element.getparent().remove(p._element)

lines = ["Integrantes del grupo (Sección A):"]
for n, (name, carnet) in enumerate(MEMBERS, 1):
    lines.append(f"{n}. {name}    Carné: {carnet}")

# Anchor: first "1. Origen" paragraph
anchor = None
for p in doc.paragraphs:
    if p.text.strip().startswith("1. Origen y fundación"):
        anchor = p
        break

if anchor is None:
    raise SystemExit("anchor not found")

for line in reversed(lines):
    anchor.insert_paragraph_before(line)

doc.save(str(path))

doc2 = Document(str(path))
for i, p in enumerate(doc2.paragraphs[:14]):
    print(i, p.text[:85] if p.text else "(empty)")
