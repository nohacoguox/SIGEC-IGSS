# -*- coding: utf-8 -*-
"""
Actualiza el borrador del Informe II y genera:
1) Solicitud de información (desde comentarios)
2) Propuesta formal al instituto
"""
from pathlib import Path
from copy import deepcopy
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BASE = Path(r"c:\Users\estua\OneDrive\Documents\UMG\10MO\seminario\INFORMES\INFORME II")
SRC = BASE / "borrador_INFORME_II_Institucion_Elegida_EORM_Agua_de_la_Mina.docx"
OUT_INFORME = BASE / "INFORME_II_Institucion_Elegida_EORM_Agua_de_la_Mina_ACTUALIZADO.docx"
OUT_SOLICITUD = BASE / "SOLICITUD_INFORMACION_EORM_Agua_de_la_Mina.docx"
OUT_PROPUESTA = BASE / "PROPUESTA_FORMAL_EORM_Agua_de_la_Mina.docx"


def set_run_font(run, name="Times New Roman", size=11, bold=False):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.bold = bold


def add_heading_styled(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        set_run_font(run, size=14 if level == 1 else 12, bold=True)
    return p


def add_para(doc, text, bold=False, size=11, align="justify", space_after=6):
    p = doc.add_paragraph()
    if align == "justify":
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    elif align == "center":
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif align == "right":
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.line_spacing = 1.15
    return p


def add_bullet(doc, text, size=11):
    p = doc.add_paragraph(style="List Bullet")
    p.clear()
    run = p.add_run(text)
    set_run_font(run, size=size)
    p.paragraph_format.space_after = Pt(3)
    return p


def add_numbered(doc, text, size=11):
    p = doc.add_paragraph(style="List Number")
    p.clear()
    run = p.add_run(text)
    set_run_font(run, size=size)
    p.paragraph_format.space_after = Pt(3)
    return p


def set_cell_text(cell, text, bold=False, size=10):
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold)


def shade_cell(cell, hex_color="1F4E79"):
    tc = cell._tePr if hasattr(cell, "_tePr") else cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_color)
    shd.set(qn("w:val"), "clear")
    tcPr.append(shd)
    for p in cell.paragraphs:
        for run in p.runs:
            run.font.color.rgb = RGBColor(255, 255, 255)


def replace_in_paragraphs(doc, old, new):
    """Replace text in paragraphs and table cells."""
    count = 0
    for p in doc.paragraphs:
        if old in p.text:
            # rebuild runs carefully if single-run or full-text match
            full = p.text
            if old in full:
                # Clear and rewrite preserving first run style if possible
                style_run = p.runs[0] if p.runs else None
                new_text = full.replace(old, new)
                for r in p.runs:
                    r.text = ""
                if p.runs:
                    p.runs[0].text = new_text
                else:
                    p.add_run(new_text)
                count += 1
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    if old in p.text:
                        full = p.text
                        new_text = full.replace(old, new)
                        for r in p.runs:
                            r.text = ""
                        if p.runs:
                            p.runs[0].text = new_text
                        else:
                            p.add_run(new_text)
                        count += 1
    return count


def replace_paragraph_containing(doc, contains, new_full_text):
    """Replace entire paragraph text if it contains a marker substring."""
    count = 0
    for p in doc.paragraphs:
        if contains in p.text:
            for r in p.runs:
                r.text = ""
            if p.runs:
                p.runs[0].text = new_full_text
            else:
                run = p.add_run(new_full_text)
                set_run_font(run)
            count += 1
    return count


def fill_blank_after_label(doc, label_substr, fill_text):
    """
    Find a paragraph containing label_substr; if following paras are blank lines
    (underscores), replace the first underscore paragraph with fill_text.
    Also handles same-paragraph cases.
    """
    paras = doc.paragraphs
    for i, p in enumerate(paras):
        if label_substr in p.text:
            # If same para has underscores after label
            if "___" in p.text:
                # leave label, hard — replace next blank instead
                pass
            # Look ahead for underscore lines
            for j in range(i + 1, min(i + 5, len(paras))):
                t = paras[j].text.strip()
                if t.startswith("___") or t == "":
                    if t.startswith("___"):
                        for r in paras[j].runs:
                            r.text = ""
                        if paras[j].runs:
                            paras[j].runs[0].text = fill_text
                        else:
                            paras[j].add_run(fill_text)
                        return True
                elif t.startswith("Ejemplo"):
                    continue
                else:
                    break
    return False


# ---------------------------------------------------------------------------
# 1) ACTUALIZAR INFORME
# ---------------------------------------------------------------------------
def update_informe():
    doc = Document(str(SRC))

    # 4.7 - Replace the paragraph about technology/internet validation
    old_47 = (
        "El registro de asistencia se realiza actualmente en papel / cuadernillos grandes. "
        "El proceso es manual, lento y difícil de consultar o actualizar por el volumen de alumnos (~320). "
        "No cuentan con dispositivos biométricos en operación. Disponen de 53 computadoras; "
        "el estado real de uso, internet y red local debe validarse in situ."
    )
    new_47 = (
        "El registro de asistencia se realiza actualmente en papel / cuadernillos grandes. "
        "El proceso es manual, lento y difícil de consultar o actualizar por el volumen de alumnos (~320). "
        "No cuentan con dispositivos biométricos en operación. Disponen de 53 computadoras. "
        "Según lo confirmado en la primera reunión formal con la dirección (31/07/2026), "
        "el establecimiento SÍ cuenta con conexión a internet y con un data center que distribuye "
        "la conectividad hacia las computadoras del laboratorio. Asimismo, se dispone de "
        "conexión eléctrica en el establecimiento. Queda pendiente detallar el uso diario exacto "
        "de las 53 computadoras y los puntos estratégicos definitivos para la instalación de los relojes biométricos."
    )
    n = replace_in_paragraphs(doc, old_47, new_47)
    print(f"4.7 main para replaced: {n}")

    # Fill internet blank
    internet_fill = (
        "Confirmado en reunión del 31/07/2026: el establecimiento cuenta con conexión a internet "
        "y un data center que distribuye la conectividad a las computadoras del laboratorio. "
        "También disponen de conexión eléctrica. Las 53 computadoras se concentran principalmente "
        "en el laboratorio; el detalle de uso diario por área queda por complementar con la dirección."
    )
    ok = fill_blank_after_label(doc, "Internet / red / uso real de las 53 computadoras", internet_fill)
    print(f"Internet blank filled: {ok}")

    # Crisis / temblor - fill blank under crisis section if present
    crisis_fill = (
        "Según lo indicado por la Directora en la reunión del 31/07/2026, durante el temblor "
        "de julio de 2025 el establecimiento sufrió daños en su infraestructura. "
        "Se solicita al instituto ampliar el detalle (áreas afectadas, estado actual de reparación "
        "y si ello condiciona la ubicación de los relojes biométricos)."
    )
    ok = fill_blank_after_label(doc, "d) Crisis o reestructuración", crisis_fill)
    if not ok:
        ok = fill_blank_after_label(doc, "Crisis o reestructuración", crisis_fill)
    if not ok:
        # Insert after any paragraph that mentions crisis heading text
        for i, p in enumerate(doc.paragraphs):
            if "Crisis o reestructuración" in p.text or "crisis o reestructuración" in p.text.lower():
                # try next few underscore paras
                for j in range(i + 1, min(i + 8, len(doc.paragraphs))):
                    t = doc.paragraphs[j].text.strip()
                    if t.startswith("___"):
                        for r in doc.paragraphs[j].runs:
                            r.text = ""
                        if doc.paragraphs[j].runs:
                            doc.paragraphs[j].runs[0].text = crisis_fill
                        else:
                            doc.paragraphs[j].add_run(crisis_fill)
                        ok = True
                        break
                break
    print(f"Crisis blank filled: {ok}")

    # 5.2 update bullet about infrastructure
    old_52 = (
        "Aprovechar la infraestructura existente (53 computadoras) e identificar condiciones "
        "de internet, energía y espacio físico para instalación."
    )
    new_52 = (
        "Aprovechar la infraestructura existente (53 computadoras, conexión a internet, data center "
        "del laboratorio y energía eléctrica) y definir, con la dirección, el espacio físico "
        "protegido (bajo techo, no a la intemperie) para instalar dos relojes biométricos en "
        "puntos estratégicos de acceso para el alumnado."
    )
    n = replace_in_paragraphs(doc, old_52, new_52)
    print(f"5.2 bullet replaced: {n}")

    # 6 intro - 2 biometric clocks
    old_6 = (
        "El grupo de seminario propone donar dispositivos biométricos e implementar un software "
        "de control de asistencia estudiantil, como aporte tecnológico de bajo costo orientado "
        "a resolver el cuello de botella del registro manual."
    )
    new_6 = (
        "El grupo de seminario propone donar e instalar dos (2) relojes biométricos e implementar "
        "un software de control de asistencia estudiantil, como aporte tecnológico orientado a "
        "resolver el cuello de botella del registro manual. Se establecen dos dispositivos por la "
        "cantidad aproximada de 320 alumnos, a fin de evitar aglomeraciones en el ingreso. "
        "El alcance del aporte se limita a lo propuesto: puesta en funcionamiento de los biométricos "
        "y del software asociado; no incluye mejoras adicionales fuera de dicho alcance."
    )
    n = replace_in_paragraphs(doc, old_6, new_6)
    print(f"6 intro replaced: {n}")

    # 6.2 first bullet
    old_62 = (
        "Dispositivos biométricos (reconocimiento facial y/o lector de huella), según "
        "factibilidad técnica, ética y de consentimiento."
    )
    new_62 = (
        "Dos (2) relojes biométricos (reconocimiento facial y/o lector de huella), con opción "
        "de código del alumno como respaldo, según factibilidad técnica, ética y de consentimiento."
    )
    n = replace_in_paragraphs(doc, old_62, new_62)
    print(f"6.2 devices replaced: {n}")

    # 6.4 considerations - expand responsibilities
    old_64 = (
        "Dado que se trata de menores de edad, el proyecto contemplará la gestión de consentimiento "
        "informado de padres o encargados cuando se utilicen datos biométricos, así como el uso de "
        "alternativas no biométricas (código del alumno) si la institución o la normativa aplicable "
        "así lo requieren. Asimismo, se validará in situ la disponibilidad eléctrica, de red y de "
        "espacio antes de fijar el alcance mínimo viable."
    )
    new_64 = (
        "Dado que se trata de menores de edad, el proyecto contemplará la gestión de consentimiento "
        "informado de padres o encargados cuando se utilicen datos biométricos, así como el uso de "
        "alternativas no biométricas (código del alumno) si la institución o la normativa aplicable "
        "así lo requieren. Se recomienda programar una reunión con padres de familia para explicar "
        "que la información biométrica recaudada será de uso exclusivo del instituto, con fines "
        "educativos y de seguridad estudiantil, para garantizar el control de asistencia. "
        "Respecto a la instalación: el grupo se encargará de dejar operativos los relojes biométricos, "
        "incluyendo las conexiones eléctricas y de red necesarias para su funcionamiento. "
        "La institución deberá designar y proporcionar la ubicación física apropiada (bajo techo, "
        "protegida de la intemperie y de la humedad/agua), en puntos estratégicos accesibles para "
        "todo el alumnado. Ya se confirmó que el establecimiento cuenta con internet, data center "
        "y energía eléctrica; lo pendiente es la designación formal del espacio de instalación."
    )
    n = replace_in_paragraphs(doc, old_64, new_64)
    print(f"6.4 replaced: {n}")

    # 6.5 next steps - meeting already done
    replacements_65 = [
        (
            "Agendar y realizar entrevista formal con la directora Sabrina Nereida Trujillo Santos de Sánchez.",
            "Primera reunión formal ya realizada el viernes 31 de julio de 2026 (14:00–15:30) con la "
            "directora Sabrina Nereida Trujillo Santos de Sánchez; continuar el seguimiento institucional.",
        ),
        (
            "Confirmar código MINEDUC, misión/visión, organigrama y estado real de infraestructura.",
            "Completar solicitud formal de información pendiente (código MINEDUC, misión/visión, "
            "organigrama, etc.); infraestructura de internet/data center/energía ya confirmada.",
        ),
        (
            "Elaborar diagnóstico breve in situ y propuesta técnica-presupuestaria de bajo costo.",
            "Entregar propuesta formal al instituto (2 relojes biométricos + software) y checklist "
            "de ubicación física / responsabilidades compartidas.",
        ),
    ]
    for o, ntext in replacements_65:
        n = replace_in_paragraphs(doc, o, ntext)
        print(f"6.5 item: {n}")

    # VII Bitácora - fill table fields by replacing placeholders in tables
    table_reps = [
        ("____ / ____ / ________", "31 / 07 / 2026"),
        ("Presencial", "PRESENCIAL (confirmada)"),  # avoid checkbox unicode issues in console
    ]
    n = replace_in_paragraphs(doc, "____ / ____ / ________", "31 / 07 / 2026")
    print(f"bitacora fecha: {n}")
    # Specific modality line from the bitacora table / example
    n = 0
    for table in doc.tables:
        for row in table.rows:
            if row.cells and row.cells[0].text.strip() == "Modalidad" and len(row.cells) > 1:
                set_cell_text(row.cells[1], "Presencial (31/07/2026, 14:00-15:30)")
                n += 1
            if row.cells and row.cells[0].text.strip() == "Fecha" and len(row.cells) > 1:
                if "31" not in row.cells[1].text:
                    set_cell_text(row.cells[1], "31 / 07 / 2026")
                    n += 1
    print(f"bitacora modality/fecha cells: {n}")

    # Replace bitácora instruction intro
    old_bit = (
        "Complete esta sección después de la primera reunión presencial o virtual con la institución. "
        "Se deja el formato listo, con espacios en blanco y un ejemplo orientativo."
    )
    new_bit = (
        "A continuación se registra la bitácora / minuta de la primera reunión formal con la institución, "
        "realizada el viernes 31 de julio de 2026."
    )
    n = replace_in_paragraphs(doc, old_bit, new_bit)
    print(f"bitacora intro: {n}")

    # Fill representative / cargo / agreements in table if still blank lines
    # Replace example bitácora with real one
    old_ex = (
        "Fecha: 12/08/2026 | Modalidad: Presencial | Lugar: Dirección de la EORM Agua de la Mina JV | "
        "Representante: Sabrina Nereida Trujillo Santos de Sánchez (Directora) | Participantes del grupo: "
        "[nombres] | Objetivo: presentar el proyecto de biométricos y software de asistencia | Acuerdos: "
        "la dirección acepta recibir la propuesta formal; se agenda visita técnica el 19/08/2026 | "
        "Próximos pasos: enviar carta de solicitud y checklist de infraestructura (responsable: coordinador del grupo)."
    )
    new_ex = (
        "Fecha: 31/07/2026 | Horario: 14:00 a 15:30 | Modalidad: Presencial | "
        "Lugar: Dirección de la EORM Agua de la Mina JV | "
        "Representante: Sabrina Nereida Trujillo Santos de Sánchez (Directora) | "
        "Objetivo: presentación del grupo y del proyecto de biométricos + software de asistencia; "
        "diagnóstico preliminar de infraestructura y necesidades. | "
        "Acuerdos / hallazgos: se confirmó que NO tienen biométricos actualmente; la asistencia se lleva "
        "en cuadernillos; SÍ hay internet, data center (laboratorio) y energía eléctrica; se mencionaron "
        "daños de infraestructura por el temblor de julio 2025; se acordó avanzar con solicitud de "
        "información pendiente y propuesta formal (2 relojes biométricos + software). | "
        "Próximos pasos: entregar solicitud de información; entregar propuesta formal; definir ubicación "
        "física de los relojes; programar reunión con padres de familia sobre datos biométricos."
    )
    n = replace_in_paragraphs(doc, old_ex, new_ex)
    print(f"bitacora example->real: {n}")

    # Fill objetivo reunión blanks
    obj_fill = (
        "Se presentó el grupo de seminario (UMG), se explicó el aporte tecnológico (relojes biométricos "
        "+ software de asistencia), se validó el problema del control manual en cuadernillos, se "
        "confirmó infraestructura de internet/data center/energía, se escuchó la mención de daños por "
        "el temblor de julio 2025 y se acordó formalizar la propuesta y la solicitud de información faltante."
    )
    ok = fill_blank_after_label(doc, "Objetivo de la reunión / puntos tratados", obj_fill)
    print(f"objetivo filled: {ok}")

    # Figure 3 pie example update
    old_f3 = (
        "Reunión presencial en la dirección de la EORM Agua de la Mina JV, 12/08/2026. "
        "Participaron la directora y [nombres del grupo]."
    )
    new_f3 = (
        "Reunión presencial en la dirección de la EORM Agua de la Mina JV, viernes 31/07/2026, "
        "de 14:00 a 15:30. Participaron la directora Sabrina Nereida Trujillo Santos de Sánchez "
        "y los integrantes del grupo de seminario. [Completar nombres exactos y pegar evidencia fotográfica.]"
    )
    n = replace_in_paragraphs(doc, old_f3, new_f3)
    print(f"fig3: {n}")

    # Conclusions update
    old_conc = (
        "La solución propuesta —donación de dispositivos biométricos e implementación de software "
        "de control de asistencia— responde de forma directa a ese requerimiento, siempre bajo "
        "consideraciones éticas de protección de datos de menores y validación técnica in situ. "
        "Queda como acción inmediata agendar la primera reunión formal con la dirección, completar "
        "la bitácora correspondiente e incorporar las evidencias gráficas del trabajo grupal en la "
        "versión PDF definitiva."
    )
    new_conc = (
        "La solución propuesta —donación e instalación de dos (2) relojes biométricos e implementación "
        "de software de control de asistencia— responde de forma directa a ese requerimiento, con "
        "alcance limitado a lo ofrecido (sin mejoras adicionales fuera de la propuesta). La primera "
        "reunión formal con la dirección ya se realizó el 31/07/2026 (14:00–15:30), confirmando "
        "internet, data center, energía eléctrica y la ausencia actual de biométricos. Quedan como "
        "acciones inmediatas: completar la información pendiente mediante solicitud formal al instituto, "
        "entregar la propuesta formal, definir la ubicación física protegida de los equipos, programar "
        "reunión con padres de familia sobre el tratamiento de datos biométricos e incorporar las "
        "evidencias gráficas del trabajo grupal en la versión PDF definitiva."
    )
    n = replace_in_paragraphs(doc, old_conc, new_conc)
    print(f"conclusiones: {n}")

    # Fuentes - entrevista
    old_fuente = (
        "Representante institucional de referencia: Directora Sabrina Nereida Trujillo Santos de Sánchez "
        "(entrevista formal: pendiente de registrar)."
    )
    new_fuente = (
        "Trujillo Santos de Sánchez, S. N. (31 de julio de 2026, 14:00–15:30). Entrevista / reunión "
        "formal presencial. Directora, EORM Agua de la Mina — Jornada Vespertina, Amatitlán, Guatemala."
    )
    n = replace_in_paragraphs(doc, old_fuente, new_fuente)
    print(f"fuentes entrevista: {n}")

    # Also update the blank entrevista formal example
    old_ent_ex = (
        "Trujillo Santos de Sánchez, S. N. (12 de agosto de 2026). Entrevista personal. Directora, "
        "EORM Agua de la Mina — Jornada Vespertina, Amatitlán, Guatemala."
    )
    new_ent_ex = (
        "Trujillo Santos de Sánchez, S. N. (31 de julio de 2026, 14:00–15:30). Entrevista / reunión "
        "formal presencial. Directora, EORM Agua de la Mina — Jornada Vespertina, Amatitlán, Guatemala."
    )
    n = replace_in_paragraphs(doc, old_ent_ex, new_ent_ex)
    print(f"fuentes ejemplo: {n}")

    # Anexo C - update infrastructure item
    old_anexo = "Validación técnica: internet, energía, espacio, modalidad biométrica y consentimientos."
    new_anexo = (
        "Validación técnica restante: designación de espacio físico protegido para 2 relojes biométricos, "
        "modalidad preferida (facial/huella/código) y consentimientos de padres. "
        "(Internet, data center y energía eléctrica: ya confirmados el 31/07/2026.)"
    )
    n = replace_in_paragraphs(doc, old_anexo, new_anexo)
    print(f"anexo C: {n}")

    # Try fill table cells for representante / cargo / acuerdos
    for table in doc.tables:
        for row in table.rows:
            cells_text = [c.text.strip() for c in row.cells]
            if not cells_text:
                continue
            label = cells_text[0]
            if label == "Lugar / enlace" and len(row.cells) > 1:
                if "____" in row.cells[1].text or row.cells[1].text.strip().startswith("___") or row.cells[1].text.strip() == "":
                    set_cell_text(row.cells[1], "Dirección, EORM Agua de la Mina — JV, Aldea Agua de la Mina, Amatitlán")
            elif label == "Representante institucional" and len(row.cells) > 1:
                if "____" in row.cells[1].text or not row.cells[1].text.strip() or row.cells[1].text.strip().startswith("___"):
                    set_cell_text(row.cells[1], "Sabrina Nereida Trujillo Santos de Sánchez")
            elif label == "Cargo" and len(row.cells) > 1:
                if "____" in row.cells[1].text or not row.cells[1].text.strip() or row.cells[1].text.strip().startswith("___"):
                    set_cell_text(row.cells[1], "Directora")
            elif label == "Participantes del grupo" and len(row.cells) > 1:
                if "____" in row.cells[1].text or row.cells[1].text.strip().startswith("___"):
                    set_cell_text(row.cells[1], "[Completar nombres de los integrantes presentes el 31/07/2026]")
            elif label == "Acuerdos" and len(row.cells) > 1:
                if "____" in row.cells[1].text or row.cells[1].text.strip().startswith("___"):
                    set_cell_text(
                        row.cells[1],
                        "Avanzar con solicitud de información pendiente y propuesta formal "
                        "(2 biométricos + software). Confirmados: sin biométricos actuales; "
                        "sí internet/data center/energía; daños por temblor jul-2025.",
                    )
            elif label.startswith("Próximos pasos") and len(row.cells) > 1:
                if "____" in row.cells[1].text or row.cells[1].text.strip().startswith("___"):
                    set_cell_text(
                        row.cells[1],
                        "1) Entregar solicitud de información. 2) Entregar propuesta formal. "
                        "3) Definir ubicación física de relojes. 4) Programar reunión con padres.",
                    )

    doc.save(str(OUT_INFORME))
    print(f"Saved: {OUT_INFORME}")


# ---------------------------------------------------------------------------
# 2) SOLICITUD DE INFORMACIÓN
# ---------------------------------------------------------------------------
def create_solicitud():
    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    add_para(doc, "UNIVERSIDAD MARIANO GÁLVEZ DE GUATEMALA", bold=True, size=12, align="center", space_after=0)
    add_para(doc, "Facultad de Ingeniería en Sistemas y Ciencias de la Computación", size=11, align="center", space_after=0)
    add_para(doc, "Seminario de Tecnología e Información", size=11, align="center", space_after=0)
    add_para(doc, "MA. Ing. Edgar Civil", size=11, align="center", space_after=12)

    add_para(doc, "SOLICITUD FORMAL DE INFORMACIÓN", bold=True, size=14, align="center", space_after=4)
    add_para(
        doc,
        "Para complementar el Informe No. II y la implementación del proyecto de automatización",
        size=11,
        align="center",
        space_after=12,
    )

    add_para(doc, "Dirigido a:", bold=True, size=11, align="justify", space_after=0)
    add_para(doc, "Profa. Sabrina Nereida Trujillo Santos de Sánchez", size=11, space_after=0)
    add_para(doc, "Directora", size=11, space_after=0)
    add_para(doc, "EORM Agua de la Mina — Jornada Vespertina (JV)", size=11, space_after=0)
    add_para(doc, "Aldea Agua de la Mina, Amatitlán, Guatemala", size=11, space_after=12)

    add_para(
        doc,
        "Respetable Directora:",
        size=11,
        space_after=8,
    )
    add_para(
        doc,
        "Por este medio, el grupo de seminario de la Universidad Mariano Gálvez de Guatemala "
        "solicita de la manera más atenta la siguiente información, con el fin de complementar "
        "el expediente académico del Informe No. II y viabilizar la implementación del aporte "
        "tecnológico (dos relojes biométricos y software de control de asistencia estudiantil).",
        size=11,
        space_after=8,
    )
    add_para(
        doc,
        "Esta solicitud se elabora a partir de los puntos pendientes identificados en el borrador "
        "del informe y de lo conversado en la primera reunión formal del viernes 31 de julio de 2026 "
        "(14:00 a 15:30).",
        size=11,
        space_after=12,
    )

    # A. Origen
    add_heading_styled(doc, "A. Origen, fundación y formalización", level=2)
    add_numbered(doc, "Código MINEDUC del establecimiento (Jornada Vespertina).")
    add_numbered(doc, "Año exacto de fundación de la Jornada Vespertina (si se conoce con precisión).")
    add_numbered(doc, "Quiénes impulsaron la creación (fundadores, comunidad, MINEDUC, vecinos, iglesia u otros actores).")
    add_numbered(doc, "¿Ha cambiado de nombre o figura legal a lo largo del tiempo? Indicar nombre(s) anteriores si aplica.")

    # B. Misión
    add_heading_styled(doc, "B. Misión y propósito institucional", level=2)
    add_numbered(doc, "Misión actual (texto oficial o declarado por la dirección).")
    add_numbered(doc, "Visión actual (si existe por escrito).")
    add_numbered(doc, "Breve descripción del propósito inicial y si se ha mantenido, ampliado o reorientado.")

    # C. Historia / crisis
    add_heading_styled(doc, "C. Historia institucional e hitos", level=2)
    add_numbered(
        doc,
        "Confirmación de la dirección actual y antigüedad en el cargo. Aclarar, si es posible, "
        "la relación con menciones previas de liderazgo (por ejemplo, referencias del año 2020 "
        "a Priscila Trujillo de Sánchez): ¿misma persona, cambio de dirección u otra situación?",
    )
    add_numbered(
        doc,
        "Detalle sobre la crisis / afectación de infraestructura por el temblor de julio de 2025 "
        "(mencionado por la dirección en la reunión del 31/07/2026): áreas afectadas, estado actual "
        "de reparación y si condiciona la ubicación de equipos.",
    )
    add_numbered(doc, "Otros hitos relevantes que la dirección considere importantes (programas, convenios, reconocimientos).")

    # D. Organización
    add_heading_styled(doc, "D. Organización y funcionamiento actual", level=2)
    add_numbered(
        doc,
        "Organigrama o descripción breve de la organización actual (dirección, docentes, "
        "personal de apoyo).",
    )
    add_numbered(
        doc,
        "(Opcional) Personal administrativo, COEDUCA y participación aproximada de padres de familia.",
    )
    add_numbered(doc, "Detalle de financiamiento o apoyos adicionales al presupuesto MINEDUC (si aplica).")

    # E. Procesos y tecnología
    add_heading_styled(doc, "E. Procesos actuales y tecnología (prioridad del proyecto)", level=2)
    add_numbered(
        doc,
        "Proceso exacto de toma de asistencia hoy: ¿quién la registra, en qué momento, con qué "
        "cuadernillo o formato, y cómo se consolida?",
    )
    add_numbered(doc, "¿Cómo generan actualmente los reportes de asistencia y con qué frecuencia / para quién?")
    add_numbered(doc, "Canales de comunicación interna y con padres de familia (WhatsApp, circulares, reuniones, etc.).")
    add_numbered(
        doc,
        "¿En algún momento intentaron manejar el proceso de asistencia de otra manera "
        "(Excel, otro sistema, formatos digitales)? Si sí, ¿qué ocurrió?",
    )
    add_numbered(
        doc,
        "Persona enlace técnico/administrativo designada para el proyecto (nombre completo, cargo, "
        "teléfono y/o correo).",
    )
    add_numbered(
        doc,
        "Confirmación de puntos de instalación posibles para dos (2) relojes biométricos: "
        "área bajo techo, con acceso a energía y red, no a la intemperie, accesible para todo el alumnado.",
    )

    # F. Modalidad y consentimientos (vía propuesta / reunión padres)
    add_heading_styled(doc, "F. Modalidad técnica y consentimientos", level=2)
    add_para(
        doc,
        "Estos puntos se formalizarán también en la propuesta del grupo; se solicita de antemano "
        "la orientación de la dirección:",
        size=11,
        space_after=6,
    )
    add_numbered(
        doc,
        "Preferencia de modalidad técnica: reconocimiento facial, huella dactilar, código del alumno "
        "o combinación.",
    )
    add_numbered(
        doc,
        "Disponibilidad para programar una reunión con padres de familia, a fin de explicar el "
        "tratamiento de la información biométrica (uso exclusivo del instituto, fines educativos "
        "y de seguridad estudiantil) y gestionar el consentimiento informado.",
    )
    add_numbered(
        doc,
        "Lineamientos institucionales o requisitos previos para el enrolamiento de menores "
        "(autorización escrita u otro mecanismo).",
    )

    add_heading_styled(doc, "G. Información ya confirmada (no se vuelve a solicitar)", level=2)
    add_bullet(doc, "Matrícula aproximada: ~320 alumnos; personal docente: ~14 catedráticos.")
    add_bullet(doc, "53 computadoras en el establecimiento.")
    add_bullet(doc, "No cuentan con biométricos en operación actualmente.")
    add_bullet(doc, "Sí cuentan con conexión a internet.")
    add_bullet(doc, "Cuentan con data center que distribuye conectividad al laboratorio.")
    add_bullet(doc, "Cuentan con conexión eléctrica.")
    add_bullet(doc, "Primera reunión formal: 31/07/2026, 14:00–15:30, con la Directora.")

    add_heading_styled(doc, "H. Forma de respuesta sugerida", level=2)
    add_para(
        doc,
        "Se agradece responder por escrito (físico o digital), o bien en una siguiente reunión "
        "presencial, a más tardar en la fecha que la dirección estime conveniente. El grupo "
        "puede proveer una guía impresa para ir llenando punto por punto.",
        size=11,
        space_after=12,
    )

    add_para(doc, "Atentamente,", size=11, space_after=24)
    add_para(doc, "_________________________________", size=11, space_after=0)
    add_para(doc, "Grupo de Seminario — UMG", bold=True, size=11, space_after=0)
    add_para(doc, "Ingeniería en Sistemas y Ciencias de la Computación", size=10, space_after=12)
    add_para(doc, "Integrantes (firmar / completar):", size=11, space_after=4)
    for i in range(1, 7):
        add_para(doc, f"{i}. _________________________________    Carné: ______________", size=10, space_after=2)

    add_para(doc, "", size=11, space_after=8)
    add_para(doc, "Recibido por la institución:", bold=True, size=11, space_after=4)
    add_para(doc, "Nombre: _________________________________    Cargo: ______________", size=10, space_after=2)
    add_para(doc, "Fecha de recepción: ____ / ____ / ________    Firma: ______________", size=10, space_after=2)

    doc.save(str(OUT_SOLICITUD))
    print(f"Saved: {OUT_SOLICITUD}")


# ---------------------------------------------------------------------------
# 3) PROPUESTA FORMAL
# ---------------------------------------------------------------------------
def create_propuesta():
    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    add_para(doc, "UNIVERSIDAD MARIANO GÁLVEZ DE GUATEMALA", bold=True, size=12, align="center", space_after=0)
    add_para(doc, "Facultad de Ingeniería en Sistemas y Ciencias de la Computación", size=11, align="center", space_after=0)
    add_para(doc, "Seminario de Tecnología e Información", size=11, align="center", space_after=0)
    add_para(doc, "MA. Ing. Edgar Civil", size=11, align="center", space_after=12)

    add_para(doc, "PROPUESTA FORMAL DE APORTE TECNOLÓGICO", bold=True, size=14, align="center", space_after=4)
    add_para(
        doc,
        "Donación e instalación de dos (2) relojes biométricos e implementación de software\n"
        "de control de asistencia estudiantil",
        size=11,
        align="center",
        space_after=12,
    )

    add_para(doc, "Institución beneficiaria:", bold=True, size=11, space_after=0)
    add_para(doc, "EORM Agua de la Mina — Jornada Vespertina (JV)", size=11, space_after=0)
    add_para(doc, "Aldea Agua de la Mina, calle principal, Amatitlán, Guatemala", size=11, space_after=0)
    add_para(doc, "Directora: Profa. Sabrina Nereida Trujillo Santos de Sánchez", size=11, space_after=12)

    add_para(doc, "1. Antecedentes", bold=True, size=12, space_after=6)
    add_para(
        doc,
        "En el marco del curso Seminario de Tecnología e Información, el grupo egresando de la "
        "Universidad Mariano Gálvez de Guatemala seleccionó a la EORM Agua de la Mina — Jornada "
        "Vespertina como institución a acompañar mediante un aporte de automatización.",
        size=11,
        space_after=6,
    )
    add_para(
        doc,
        "El 31 de julio de 2026, de 14:00 a 15:30, se sostuvo la primera reunión formal con la "
        "Directora. En dicha reunión se confirmó que el control de asistencia se realiza de forma "
        "manual en cuadernillos; que no existen biométricos en operación; y que el establecimiento "
        "sí cuenta con conexión a internet, un data center que distribuye conectividad al laboratorio "
        "y energía eléctrica. La matrícula aproximada es de 320 alumnos, con alrededor de 14 catedráticos "
        "y 53 computadoras.",
        size=11,
        space_after=12,
    )

    add_para(doc, "2. Problema que se busca resolver", bold=True, size=12, space_after=6)
    add_bullet(doc, "La asistencia manual en cuadernillos consume tiempo y recursos del personal.")
    add_bullet(doc, "Es difícil disponer de datos actualizados y confiables ante ~320 alumnos.")
    add_bullet(
        doc,
        "Existen casos en que se reporta en hogares que el alumno asiste, pero no llega al "
        "establecimiento, lo que debilita el control institucional.",
    )
    add_para(doc, "", size=6, space_after=0)

    add_para(doc, "3. Objetivo de la propuesta", bold=True, size=12, space_after=6)
    add_para(
        doc,
        "Digitalizar y automatizar el control de asistencia estudiantil de la EORM Agua de la Mina — JV, "
        "mediante la donación e instalación de dos relojes biométricos y un software de gestión, "
        "mejorando la trazabilidad, oportunidad y confiabilidad de la información para dirección, "
        "docentes y familias.",
        size=11,
        space_after=12,
    )

    add_para(doc, "4. Alcance de lo que ofrece el grupo (y límite explícito)", bold=True, size=12, space_after=6)
    add_para(
        doc,
        "El aporte se limita estrictamente a lo descrito en esta propuesta. El grupo no se compromete "
        "a realizar mejoras adicionales, remodelaciones, ni desarrollos fuera del alcance aquí definido.",
        size=11,
        space_after=6,
    )
    add_para(doc, "El grupo se compromete a:", bold=True, size=11, space_after=4)
    add_bullet(
        doc,
        "Donar e instalar dos (2) relojes biométricos. Se establecen dos equipos por el volumen "
        "aproximado de 320 alumnos, para reducir aglomeraciones en el ingreso y facilitar el flujo.",
    )
    add_bullet(
        doc,
        "Dejar los biométricos funcionando, incluyendo la gestión de las conexiones eléctricas y "
        "de red necesarias para su operación (aprovechando que el instituto ya cuenta con energía, "
        "internet y data center).",
    )
    add_bullet(
        doc,
        "Desarrollar e implementar un software de control de asistencia (registro, consulta y reportes).",
    )
    add_bullet(
        doc,
        "Habilitar modalidades de marcaje según factibilidad y autorización institucional: "
        "reconocimiento facial, huella dactilar y/o código del alumno (como respaldo o alternativa).",
    )
    add_bullet(doc, "Capacitar al personal enlace designado por la dirección.")
    add_bullet(doc, "Entregar documentación básica de uso y recomendaciones de resguardo de la información.")
    add_para(doc, "", size=6, space_after=0)
    add_para(doc, "Queda fuera del alcance:", bold=True, size=11, space_after=4)
    add_bullet(doc, "Obras civiles, remodelaciones o techados adicionales (si se requirieran, corren por cuenta de la institución).")
    add_bullet(doc, "Compra de computadoras adicionales, servidores o ampliación de infraestructura no relacionada.")
    add_bullet(doc, "Cualquier mejora tecnológica o administrativa no descrita en esta propuesta.")
    add_para(doc, "", size=6, space_after=0)

    add_para(doc, "5. Lo que se requiere de la institución", bold=True, size=12, space_after=6)
    add_para(
        doc,
        "Para que el aporte sea operativamente viable, la institución deberá:",
        size=11,
        space_after=6,
    )
    add_numbered(
        doc,
        "Designar y proporcionar la ubicación física apropiada para los dos relojes biométricos. "
        "Los equipos NO pueden quedar a la intemperie: requieren protección bajo techo, no son "
        "equipos a prueba de agua y llevan conexión eléctrica.",
    )
    add_numbered(
        doc,
        "Ubicar los relojes en puntos estratégicos de acceso, de modo que todo el alumnado pueda "
        "marcar sin dificultad y se eviten cuellos de botella.",
    )
    add_numbered(
        doc,
        "Facilitar el acceso a los puntos de energía eléctrica y de red/internet necesarios en las "
        "ubicaciones designadas (la institución ya cuenta con estos servicios a nivel general).",
    )
    add_numbered(
        doc,
        "Designar una persona enlace (docente o administrativo) para coordinación, pruebas y capacitación.",
    )
    add_numbered(
        doc,
        "Apoyar la programación de una reunión con padres de familia para informar sobre la recolección "
        "de datos biométricos: su uso será exclusivamente del instituto, con fines educativos y de "
        "seguridad de los estudiantes, orientado a garantizar el registro real de asistencia.",
    )
    add_numbered(
        doc,
        "Gestionar o facilitar el consentimiento informado de padres/encargados cuando se utilicen "
        "datos biométricos de menores; quienes no autoricen podrán utilizar la modalidad de código "
        "del alumno, si así se define.",
    )
    add_numbered(
        doc,
        "Emitir, de ser posible, una carta o constancia de aceptación del aporte académico del grupo.",
    )
    add_para(doc, "", size=6, space_after=0)

    add_para(doc, "6. Consideraciones éticas y de protección de datos", bold=True, size=12, space_after=6)
    add_para(
        doc,
        "Al tratarse de menores de edad, el tratamiento de información biométrica se realizará "
        "únicamente para fines del control de asistencia educativa del establecimiento. "
        "Se promoverá la transparencia ante la comunidad educativa (reunión con padres) y se "
        "respetarán las alternativas no biométricas cuando no exista autorización.",
        size=11,
        space_after=12,
    )

    add_para(doc, "7. Beneficios esperados", bold=True, size=12, space_after=6)
    add_bullet(doc, "Reducción del tiempo dedicado al llenado y consolidación de cuadernillos.")
    add_bullet(doc, "Mayor confiabilidad y trazabilidad (fecha/hora) del registro de asistencia.")
    add_bullet(doc, "Consultas y reportes digitales más ágiles para dirección y docentes.")
    add_bullet(doc, "Mejor control institucional frente a inconsistencias reportadas desde los hogares.")
    add_bullet(doc, "Aprovechamiento de la infraestructura ya existente (internet, data center, energía, computadoras).")
    add_para(doc, "", size=6, space_after=0)

    add_para(doc, "8. Plan de trabajo tentativo", bold=True, size=12, space_after=6)
    table = doc.add_table(rows=5, cols=2)
    table.style = "Table Grid"
    headers = [("Fase", "Actividad")]
    rows = [
        ("1. Formalización", "Aceptación de la propuesta; designación de enlace; definición de ubicación de los 2 relojes."),
        ("2. Preparación", "Reunión con padres; consentimientos; relevamiento técnico de puntos de instalación."),
        ("3. Implementación", "Instalación eléctrica/red de los biométricos; despliegue del software; pruebas."),
        ("4. Cierre", "Capacitación, entrega de documentación y puesta en operación."),
    ]
    set_cell_text(table.rows[0].cells[0], "Fase", bold=True, size=10)
    set_cell_text(table.rows[0].cells[1], "Actividad", bold=True, size=10)
    for i, (a, b) in enumerate(rows, start=1):
        set_cell_text(table.rows[i].cells[0], a, bold=True, size=10)
        set_cell_text(table.rows[i].cells[1], b, size=10)
    doc.add_paragraph()

    add_para(doc, "9. Solicitud de aceptación", bold=True, size=12, space_after=6)
    add_para(
        doc,
        "Se solicita respetuosamente a la Dirección de la EORM Agua de la Mina — Jornada Vespertina "
        "aceptar la presente propuesta de aporte tecnológico, en los términos de alcance y "
        "responsabilidades aquí descritos, para continuar con las fases de implementación del "
        "proyecto de seminario.",
        size=11,
        space_after=16,
    )

    add_para(doc, "Atentamente,", size=11, space_after=24)
    add_para(doc, "_________________________________", size=11, space_after=0)
    add_para(doc, "Grupo de Seminario — UMG", bold=True, size=11, space_after=0)
    add_para(doc, "Ingeniería en Sistemas y Ciencias de la Computación", size=10, space_after=12)
    add_para(doc, "Integrantes (firmar / completar):", size=11, space_after=4)
    for i in range(1, 7):
        add_para(doc, f"{i}. _________________________________    Carné: ______________", size=10, space_after=2)

    add_para(doc, "", size=11, space_after=10)
    add_para(doc, "ACEPTACIÓN INSTITUCIONAL", bold=True, size=12, align="center", space_after=8)
    add_para(
        doc,
        "Yo, _______________________________________________, en mi calidad de "
        "_______________________________________________ de la EORM Agua de la Mina — Jornada "
        "Vespertina, manifiesto conocer el alcance de esta propuesta y:",
        size=11,
        space_after=8,
    )
    add_para(doc, "☐  ACEPTO la propuesta en los términos descritos.", size=11, space_after=4)
    add_para(doc, "☐  ACEPTO con observaciones: _______________________________________________", size=11, space_after=4)
    add_para(doc, "☐  NO ACEPTO. Motivo: _______________________________________________________", size=11, space_after=12)
    add_para(doc, "Lugar y fecha: _____________________________, ____ / ____ / ________", size=11, space_after=16)
    add_para(doc, "Firma y sello: _________________________________", size=11, space_after=2)
    add_para(doc, "Nombre: _________________________________", size=11, space_after=2)
    add_para(doc, "Cargo: _________________________________", size=11, space_after=2)

    doc.save(str(OUT_PROPUESTA))
    print(f"Saved: {OUT_PROPUESTA}")


if __name__ == "__main__":
    print("=== UPDATE INFORME ===")
    update_informe()
    print("\n=== SOLICITUD ===")
    create_solicitud()
    print("\n=== PROPUESTA ===")
    create_propuesta()
    print("\nDONE")
